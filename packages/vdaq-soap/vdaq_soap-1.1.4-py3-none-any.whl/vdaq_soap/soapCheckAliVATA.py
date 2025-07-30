#!/usr/bin/env python3
"""Create a run and execute it."""

import sys
import os
import json
import zlib
import math
import pathlib
import time
from argparse import ArgumentParser
import numpy as np
import matplotlib.pyplot as plt
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    import daqpp.soap
except ModuleNotFoundError:
    pass

from vdaq_ana import GTimer, ShowProgress

from vdaq_soap.data_utils import fit_gaussian, draw_best_fit, open_data_file
from vdaq_soap.analyzeHoldDelay import read_hold_delay_data
from vdaq_soap.analyzeMbias import read_mbias_scan


def set_module_param(runm, par_name, value):
    """Set value of a given parameter for all modules in RunManager.

    Args:
    ----
        runm: the RunManager
        par_name: The parameter name
        value: The value for the parameter

    """
    kkk = runm.getModules()
    kkk = kkk.values()
    for md in runm.getModules().values():
        try:
            md.parameters[par_name].set_value(value)
        except IndexError:
            print("parameter {} does not exist for module {}]".format(par_name, md.name))


def set_default_parameters(run_manager, hold_delay=200, polarity=0, mbias=600, threshold=50):
    """Set default module parameters."""
    # print("Def. parameters. Setting HD to {} Mbias {}".format(hold_delay, mbias))
    set_module_param(run_manager, "channel", -9999)
    set_module_param(run_manager, "coincidence", False)
    set_module_param(run_manager, "ro_mode", 4)      # sparse
    set_module_param(run_manager, "threshold", threshold)  # threshold
    set_module_param(run_manager, "polarity", polarity)
    set_module_param(run_manager, "nadj", 5)
    set_module_param(run_manager, "trg_type", 1)
    set_module_param(run_manager, "hold_delay", hold_delay)

    try:
        version = run_manager.getParameter("version")
        set_module_param(run_manager, "mbias", mbias)
    except Exception:
        set_module_param(run_manager, "mbias", [mbias, mbias, mbias, mbias])


def wait_for_end_of_run(run_manager, max_time=1e10):
    """Waits for end of run.

    It also displays the state of the run.

    Args:
    ----
        run_manager: The run_manager
        max_time: maximum run_duration

    """
    gt = GTimer()
    gt.start()
    forced = False
    while True:
        S = run_manager.getStatus()
        if S is None:
            continue

        if S.status != "Running" and S.status != "Paused":
            break

        if gt() > max_time and S.status:
            print("\nHave to stop the run")
            forced = True
            run_manager.stopRun()
            break


        print('\r', S, end=' ')
        sys.stdout.flush()
        time.sleep(0.5)

    if not forced:
        print("\n")

    time.sleep(0.5)
    S = run_manager.getStatus()
    print("End of Run: ", S)


def get_monitor_data(module):
    """Get the monitor data.

    Args:
    ----
        module: the module object

    """
    mdata = module.getMonitorData()
    mdata = zlib.decompress(mdata)
    data = json.loads(mdata)
    return data


def draw_hist_and_projection(axh, axp, data, title=None, axis_title=None, x_label=None, y_label=None):
    """Plots a bar histogram and the Y projection.

    Args:
    ----
        axh: axis for the histogram
        axp: axis for the projection
        data: The monitor data
        title: the histogram title
        axis_title: title for the histogram axis
        x_label: label for the histgram X axis
        y_label: label for the Y axis

    Returns
    -------
        mean. std: a tuple with the Y maverage and std

    """
    # Draw the histogram
    y = data['data']
    x = np.linspace(data['xmin'], data['xmax'], data['nx'])

    axh.step(x, y)
    if axis_title:
        axh.set_title(axis_title)

    if x_label:
        axh.set_xlabel(x_label)

    if y_label:
        axh.set_ylabel(y_label)

    mean = np.mean(y)
    std = np.std(y)

    # Now the projection
    axp.hist(y, orientation="horizontal")

    return mean, std


def test_pedestals(server, the_module, save_data=False, document=None):
    """Makes a pedestal run."""
    print("\n:> Testing pedestals")

    # Create a new scan run
    server.createRunManager("pedestal", "main")
    run_manager = daqpp.soap.RunManager("main", server)
    set_module_param(run_manager, "ro_mode", 1)

    # Reset all monitor data
    the_module.resetMonitorData()

    # Now we make the pedestal run
    if save_data:
        ofile = pathlib.Path.cwd().joinpath("pedestal_{}.h5".format(the_module.name))
        run_manager.getParameter("command").set_value(
            "logData|file={}".format(ofile))

    run_manager.setMaxEvents(10000)
    run_manager.startRun()
    wait_for_end_of_run(run_manager, 5.0)

    # Get the data and plot
    fig = plt.figure(constrained_layout=True)
    gs = fig.add_gridspec(ncols=2, nrows=2, width_ratios=(2, 1))

    # Get the data
    data = get_monitor_data(the_module)

    # pedestals
    ax = fig.add_subplot(gs[0, 0])
    mped, stped = draw_hist_and_projection(
        ax, fig.add_subplot(gs[0, 1], sharey=ax),
        data['histograms']['pedestal'],
        axis_title="Pedestals",
        x_label="Channel"
    )
    print("Pedestal mean {:.1f} std {:.1f}".format(mped, stped))

    # noise
    ax = fig.add_subplot(gs[1, 0])
    mnoise, stnoise = draw_hist_and_projection(
        ax, fig.add_subplot(gs[1, 1], sharey=ax),
        data['histograms']['noise'],
        axis_title="Noise",
        x_label="Channel"
    )
    print("Noise mean {:.1f} std {:.1f}".format(mnoise, stnoise))

    pdf_file = pathlib.Path.cwd().joinpath("pedestals_{}.png".format(the_module.name))
    fig.savefig(pdf_file, dpi=300)

    if document:
        document.add_heading('Pedestals', level=2)
        p = document.add_paragraph("Pedestal mean {:.1f} std {:.1f}".format(mped, stped))
        p = document.add_paragraph("Noise mean {:.1f} std {:.1f}".format(mnoise, stnoise))
        document.add_picture(str(pdf_file))
        os.remove(pdf_file)

    plt.draw()
    plt.pause(0.001)


def test_ext_trigger(server, the_module, nsec=5, document=False):
    """Test the external trigger on the given module.

    Args:
    ----
        server: the SOAP client
        the_module: the module
        nsec: duration of the run. Default is 5 sec.
        document: the word document

    """
    print("\n:> Testing Ext Trigger")

    # Now we move back to a data run
    server.createRunManager("data", "main")
    run_manager = daqpp.soap.RunManager("main", server)

    old_coinc = the_module.parameters['coincidence'].get_value()
    old_trg_type = the_module.parameters['trg_type'].get_value()
    old_ro_mode = the_module.parameters['ro_mode'].get_value()

    # Set the module trigger mode to Ext.
    set_module_param(run_manager, "ro_mode", 1)
    set_module_param(run_manager, "trg_type", 4)
    set_module_param(run_manager, "coincidence", True)

    # Reset all monitor data
    the_module.resetMonitorData()

    # Now we start the 5 sec. run
    run_manager.setMaxEvents(-1)
    run_manager.startRun()
    wait_for_end_of_run(run_manager, nsec)

    S = run_manager.getStatus()
    print("Got {} with a rate of {:.1f} Hz with Ext. trigger.".format(S.ntrigger, S.rate))
    if document:
        document.add_heading('Ext. Trigger', level=2)
        p = document.add_paragraph("Got {} events with a rate of {:.1f} Hz.".format(S.ntrigger, S.rate))

    # Restore the values
    the_module.parameters['trg_type'].set_value(old_trg_type)
    the_module.parameters['ro_mode'].set_value(old_ro_mode)
    the_module.parameters['coincidence'].set_value(old_coinc)


def test_hold_delay_scan(server, the_module, chan, npts=32, first=0, last=255, nsec=1, emin=-1, document=None):
    """Make a HoldDelay scan.

    Args:
    ----
        server: The SOAP server
        the_module: The module to test
        chan: The channel tto send teh calib pulse to.
        npts: number of points.
        first: First hold_delay value. Defaults to 0.
        last: Last hold_delay value. Defaults to 255.
        nsec: number of seconds per scan point
        emin (optional): min value to store in histrogram.
        document: the word document

    """
    print("\n:> Hold delay scan (chan {}: {}, {}, {}".format(chan, first, last, npts))

    # Create a new scan run
    server.createRunManager("scan", "main")
    run_manager = daqpp.soap.RunManager("main", server)

    # Reset all monitor data
    the_module.resetMonitorData()

    # Get current hold_delay
    old_hold_delay = the_module.parameters['hold_delay'].get_value()

    # Recompute scan range
    step = math.floor((last-first)/(npts-1))
    last = first + (npts-1) * step
    if last > 255:
        last -= step

    # Define the scan
    run_manager.getParameter("command").set_value(
        "scan|nevt={}:is_time=true:points=(1,1,0,0);(0,1,{},{});(11,{},{},{})".format(
            nsec, chan, chan, npts, first, last)
    )

    # Set the output file
    ofile = pathlib.Path.cwd().joinpath("hold_delay_{}.h5".format(the_module.name))
    run_manager.getParameter("command").set_value(
        "logData|file={}".format(ofile))

    # Start the run
    run_manager.startRun()
    wait_for_end_of_run(run_manager)
    time.sleep(0.5)

    # restore hold_delay value
    the_module.parameters['hold_delay'].set_value(old_hold_delay)

    # read the data
    ofile = pathlib.Path.cwd().joinpath("hold_delay_{}_000.h5".format(the_module.name))
    hd_max = read_hold_delay_data(ofile, int(the_module.name), emin=emin, document=document)
    return hd_max


def test_channel_change(server, the_module, nchan, first_chan, last_chan, nevts=1000, document=None):
    """Make a channel scan.

    Args:
    ----
        server: the SOAP client
        the_module: The module to test
        nchan: number of channels
        first_chan: first channel
        last_chan: Last channel
        nevts (optional): Number of events per channel. Defaults to 1000.
        document (optional): the word document

    """
    print("\n:> Channel scan (chan {} {}-{}".format(nchan, first_chan, last_chan))

    # Create a new scan run
    server.createRunManager("scan", "main")
    run_manager = daqpp.soap.RunManager("main", server)

    # Reset all monitor data
    the_module.resetMonitorData()

    # Recompute scan range
    if nchan > 1:
        nchip = the_module.parameters['nchip'].get_value()
        ntot = nchip * 128
        step = math.floor((last_chan-first_chan)/(nchan-1))
        last_chan = first_chan + (nchan-1) * step
        if last_chan > ntot:
            last_chan -= step
    else:
        if nchan <= 0:
            nchan = 1
        step = 1
        last_chan = first_chan

    print("New scan range: {} {}-{} [{}]".format(nchan, first_chan, last_chan, step))

    # Define the scan
    run_manager.getParameter("command").set_value(
        "scan|nevt={}:is_time=false:points=(1,1,0,0);(0,{},{},{})".format(
            nevts, nchan, first_chan, last_chan)
    )

    # Set the output file
    ofile = pathlib.Path.cwd().joinpath("channel_scan_{}.h5".format(the_module.name))
    run_manager.getParameter("command").set_value(
        "logData|file={}".format(ofile))

    run_manager.startRun()
    wait_for_end_of_run(run_manager)
    run_status = run_manager.getStatus()

    # Get the data and check the hitmap.
    data = get_monitor_data(the_module)
    hitmap = np.array(data['histograms']['hitmap']['data'])

    channels = [x for x in range(first_chan, last_chan, nchan) if x < ntot]
    if len(channels) > 1:
        values = (hitmap[channels] - nevts)/nevts
        mean = np.mean(values)
        std = np.std(values)
        print("Hitmap mean {:.1f} std {:.1f}. Expected: {}".format(mean, std, 0.0))
        print(channels)
        print(values)

        # Draw
        fig, ax = plt.subplots(1, 1)
        PH = data['histograms']['hitmap']
        xin = np.linspace(PH['xmin'], PH['xmax'], PH['nx'])

        ax.step(xin, hitmap)
        ax.set_title("Hitmap")
        ax.set_xlabel("Channel")

        pdf_file = pathlib.Path.cwd().joinpath("channels_{}.png".format(the_module.name))
        fig.savefig(pdf_file, dpi=300)
        if document:
            document.add_heading('Channel Scan', level=2)
            p = document.add_paragraph("Expected number of events per channel: {}.".format(nevts))
            p = document.add_paragraph("(Hitmap-nevt)/nevt mean {:.1f} std {:.1f}. Expected: {}.".format(mean,
                                                                                                         std,
                                                                                                         0.0))
            document.add_picture(str(pdf_file))
            os.remove(pdf_file)

        plt.draw()
        plt.pause(0.001)

    return run_status


def test_TDCs(server, the_module, channel, nsec=5, emin=0.0, document=None):
    """Test the TDC or clock counter.

    Args:
    ----
        server: The soap client
        the_module: the Module
        channel: the channel
        nsec (optional): Acquisition time. Defaults to 5 sec.
        emin (optional): min value to store in histrogram.
        document (optional): The report document. Defaults to None.

    """
    # do a channel scan for this
    ofile = pathlib.Path.cwd().joinpath("channel_TDC_{}.h5".format(the_module.name))
    server.createRunManager("data", "main")
    run_manager = daqpp.soap.RunManager("main", server)
    set_module_param(run_manager, "channel", channel)
    run_manager.GetReady()
    run_manager.getParameter("command").set_value("logData|file={}".format(ofile))
    # Reset all monitor data
    the_module.resetMonitorData()

    # Take data
    run_manager.startRun()
    wait_for_end_of_run(run_manager, nsec)
    status = run_manager.getStatus()
    rate = status.rate

    # Now we have to read back the data
    ifile = pathlib.Path.cwd().joinpath("channel_TDC_{}_000.h5".format(the_module.name))

    print(":> Opening {} for TDC analysis".format(ifile))
    vdaq = open_data_file(ifile)

    if vdaq is None:
        return

    last_time = {}
    last_event = {}
    nlost = 0
    mod_id = int(the_module.name)
    values = []
    amplitude = []

    prg = ShowProgress(vdaq.nevts, width=24)
    prg.start()
    for evt in vdaq:
        evt_time = int(evt.time)
        mid = evt.mod_id
        if mid != mod_id:
            continue

        ltim = last_time.get(mid, -1)
        if ltim > 0:
            dt = evt_time - ltim
        else:
            dt = 0

        last_time[mid] = evt_time

        levent = last_event.get(mid, -1)
        if levent > 0:
            devt = evt.evtcnt - levent
            if devt > 1:
                nlost += 1

        last_event[mid] = evt.evtcnt
        dt = (dt*25.0)/1000.0
        if dt > 0 and dt < 1000:
            values.append(dt)

        # The single channel amplitude
        md = vdaq.modules[mid]
        data = md.process_event(evt)
        if data is not None:
            for C, E in data:
                if E > emin:
                    amplitude.append(E)

        prg.increase(show=True)

    print("\nNumber of events lost: {}".format(nlost))
    if rate > 0:
        print("Expected rate: {:.2f} Hz. Period {:.6f} us.".format(rate, 1.0e6/rate))
    else:
        print("Could not get rate from RM status: {}".format(status))

    print("Mean period: {:.3f} us".format(np.mean(values)))

    fig, ax = plt.subplots(1, 1)
    ax.hist(values, bins=50)
    ax.set_title("TDC test")
    ax.set_xlabel("Period (us)")

    pdf_file = pathlib.Path.cwd().joinpath("tdc_{}.png".format(the_module.name))
    fig.savefig(pdf_file, dpi=300)
    if document:
        document.add_heading('Test of TDC values', level=2)
        p = document.add_paragraph("Number of events lost: {}".format(nlost))
        if rate > 0.0:
            p = document.add_paragraph("Expected rate: {:.2f} Hz. Period {:.6f} us.".format(rate, 1.0e6/rate))
        else:
            p = document.add_paragraph("Could not get the RM status.")

        P = document.add_paragraph("Mean period: {:.3f} us".format(np.mean(values)))
        document.add_picture(str(pdf_file))
        os.remove(pdf_file)

    # Draw the signal
    fig, ax = plt.subplots(1, 1)
    n, bins, *_ = ax.hist(amplitude, bins=50)
    step = 0.5 * (bins[1] - bins[0])
    X = bins[:-1] + step
    mean = np.mean(amplitude)
    std = np.std(amplitude)
    try:
        result, legend = fit_gaussian(n, X, mean, width=std)
        draw_best_fit(ax, result, bins)
        ax.legend([legend], loc=1)
        ax.set_title("Signel Channel signal")
        ax.set_xlabel("Charge (ADC)")
        pdf_file = pathlib.Path.cwd().joinpath("tdc_{}.png".format(the_module.name))
        fig.savefig(pdf_file, dpi=300)
        if document:
            document.add_heading('Single Channel Spectrum', level=2)
            P = document.add_paragraph("Mean : {:.3f} Std {:3f}".format(mean, std))
            document.add_picture(str(pdf_file))
            os.remove(pdf_file)

    except Exception as ex:
        print("something went wrong woth TDC. {}".format(str(ex)))

    plt.draw()
    plt.pause(0.001)


def test_FOR_run(server, nevts):
    """Test the FOR trigger.

    Args:
    ----
        server: The SOAP client
        nevts: Number of events

    """
    print("\n:> Normal FOR run")
    server.createRunManager("data", "main")
    run_manager = daqpp.soap.RunManager("main", server)
    modules = server.getAllModules()
    for md in modules:
        md.parameters['channel'].set_value(-1)
        md.resetMonitorData()
        md.setLocal(False, "main")

    run_manager.GetReady()
    run_manager.startRun()
    wait_for_end_of_run(run_manager, 6)
    print("Done")


def test_mbias_scan(server, the_module,
                    nchan, first_chan, last_chan,
                    npts, first, last,
                    nsec=1, emin=-1, document=None):
    """Make a mbias scan.

    Args:
    ----
        server: the SOAP client
        the_module: The module to test
        nchan: number of channels
        first_chan: First chan in range
        last_chan: last chan in range
        npts: number of mbias points
        first: first mbias value
        last: last mbias value
        nsec: Tiem on each scanpoint. Defaults to 1.
        document (optional): the word document

    """
    print("\n:> mbias scan (chan {} {}-{}, sec {}): mbias from {} to {} npts {}".format(nchan, first_chan, last_chan, nsec, first, last, npts))

    # Create a new scan run
    server.createRunManager("scan", "main")
    run_manager = daqpp.soap.RunManager("main", server)

    # Reset all monitor data
    the_module.resetMonitorData()

    # Recompute scan range
    # mbias
    if npts > 1:
        step = math.floor((last-first)/(npts-1))
        last = first + (npts-1) * step
    else:
        last = first

    # channel
    nchip = the_module.parameters['nchip'].get_value()
    ntot = 128*nchip
    if nchan > 1:
        step = math.floor((last_chan-first_chan)/(nchan-1))
        last_chan = first_chan + (nchan-1) * step
    else:
        last_chan = first_chan

    # Define the scan
    if nchan == 1 and first_chan < 0:
        run_manager.getParameter("command").set_value(
            "scan|nevt={}:is_time=true:points=(4,{},{},{})".format(
                nsec, npts, first, last)
        )
    else:
        run_manager.getParameter("command").set_value(
            "scan|nevt={}:is_time=true:points=(1,1,0,0);(0,{},{},{});(4,{},{},{})".format(
                nsec, nchan, first_chan, last_chan, npts, first, last)
        )

    # Set the output file name
    ofile = pathlib.Path.cwd().joinpath("mbias_{}.h5".format(the_module.name))
    run_manager.getParameter("command").set_value("logData|file={}".format(ofile))

    # start the run
    run_manager.startRun()
    wait_for_end_of_run(run_manager)

    # Wait to have data file closed
    time.sleep(0.5)

    # read the data
    ofile = pathlib.Path.cwd().joinpath("mbias_{}_000.h5".format(the_module.name))
    read_mbias_scan(ofile, int(the_module.name), emin=emin, document=document)


def restore_default_state(server, options):
    """Restore state and parameters."""
    # BAck to normal Data RunManager.
    server.createRunManager("data", "main")
    run_manager = daqpp.soap.RunManager("main", server)

    # Set all settings in the default state
    set_default_parameters(run_manager, hold_delay=options.hold_delay, polarity=options.polarity, mbias=options.mbias)


def soapCheckAliVATA(options):
    """Connect to server and start a scan run."""
    # This is where we connect with the server
    try:
        server = daqpp.soap.DAQppClient(options.host, options.port, debug=options.debug)
    except Exception as E:
        print(("Could not connect to the server\n{}".format(E.response)))
        return

    modid = options.mid
    channel = options.channel
    document = docx.Document()
    if options.title:
        document.add_heading(options.title, 0)

    if options.firmware:
        P = document.add_heading("Firmware: {}".format(options.firmware), level=1)
        P.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading('Slot {}'.format(modid), level=1)
    if options.pulse is not None:
        document.add_paragraph("Calibration pulse amplitude: {:.2f} mV.".format(options.pulse))

    # Stop the run if running
    print("Check if the RUnManager is active and running")
    run_manager = daqpp.soap.RunManager("main", server)
    S = run_manager.getStatus()
    if S.status == "Running":
        print("Stopping run")
        run_manager.stopRun()

    # This is the slot or module id we aht to check.
    in_local = []

    # Set the other modules in local (out of main run manager)
    the_module = None
    modules = server.getAllModules()
    for md in modules:
        print("Module {}".format(md.name))
        if md.name != modid:
            md.setLocal(True, "")
            in_local.append(md)
        else:
            md.setLocal(False, "main")
            the_module = md

    if the_module is None:
        print("### CheckAliVATA Error: module {} not present")
        return

    # Set some module parameters like adjacent, polarity and threshold.
    set_default_parameters(run_manager, hold_delay=options.hold_delay, polarity=options.polarity, mbias=options.mbias, threshold=options.threshold)
    run_manager.GetReady()

    # Test the pedestals
    set_default_parameters(run_manager, hold_delay=options.hold_delay, polarity=options.polarity, mbias=options.mbias, threshold=options.threshold)
    test_pedestals(server, the_module, save_data=True, document=document)

    # Test Ext Trigger
    set_default_parameters(run_manager, hold_delay=options.hold_delay, polarity=options.polarity, mbias=options.mbias, threshold=options.threshold)
    test_ext_trigger(server, the_module, nsec=5, document=document)

    # Do a channel scan
    set_default_parameters(run_manager, hold_delay=options.hold_delay, polarity=options.polarity, mbias=options.mbias, threshold=options.threshold)
    nchan = options.nchip*8
    first_chan = 8
    last_chan = options.nchip*128 - 8
    test_channel_change(server, the_module, nchan, first_chan, last_chan, 5000, document=document)

    # test TDCs
    set_default_parameters(run_manager, hold_delay=options.hold_delay, polarity=options.polarity, mbias=options.mbias, threshold=options.threshold)
    test_TDCs(server, the_module, channel, nsec=5, emin=options.emin, document=document)

    # Do a hold delay scan
    set_default_parameters(run_manager, hold_delay=options.hold_delay, polarity=options.polarity, mbias=options.mbias, threshold=options.threshold)
    test_hold_delay_scan(server, the_module, channel, 64, 0, 255, nsec=1, emin=options.emin, document=document)

    # Do a mbias scan
    set_default_parameters(run_manager, hold_delay=options.hold_delay, polarity=options.polarity, mbias=options.mbias, threshold=options.threshold)
    test_mbias_scan(server, the_module,
                    1, channel, channel,
                    64, 100, 4096,
                    nsec=1, emin=options.emin, document=document)

    # Now we move back to a data run
    # test_FOR_run(server, 10000)

    # Set all settings in the default state
    restore_default_state(server, options)
    document.save(options.out)


def main():
    """Main entry."""
    parser = ArgumentParser()
    parser.add_argument("--mid", dest="mid", default="11",
                        type=str, help="The slot ID")
    parser.add_argument("--out", dest="out", default="vdaq_report.docx",
                        type=str, help="The output fiel name")
    parser.add_argument("--title", dest="title", default=None,
                        type=str, help="Document title")
    parser.add_argument("--firmware", dest="firmware", default="XX.XX.XX",
                        type=str, help="Firmware")
    parser.add_argument("--host", dest="host", default="localhost",
                        type=str, help="The soap server")
    parser.add_argument("--port", dest="port", default=50000,
                        type=int, help="The soap port")
    parser.add_argument("--channel", dest="channel", default=32,
                        type=int, help="The test channel")
    parser.add_argument("--pulse", dest="pulse", default=None,
                        type=float, help="Cal. pulse amplitude")
    parser.add_argument("--mbias", dest="mbias", default=300,
                        type=int, help="Chip Main bias")
    parser.add_argument("--hold_delay", dest="hold_delay", default=175,
                        type=int, help="Chip Hold delay")
    parser.add_argument("--polarity", dest="polarity", default=0,
                        type=int, help="Signal polarity")
    parser.add_argument("--threshold", default=50, type=int, help="GPx threshold")
    
    parser.add_argument("--emin", default=0.0, type=float, help="Min E to show in histogram")
    parser.add_argument("--nchip", help="Number of chips", default=1, type=int)
    parser.add_argument("--debug", dest="debug", action="store_true",
                        help="Debug server I/O",
                        default=False)

    options = parser.parse_args()

    plt.ion()

    soapCheckAliVATA(options)

    print("\n### All tests done ###")
    plt.ioff()
    # plt.show()


if __name__ == "__main__":
    main()
