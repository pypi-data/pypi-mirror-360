
import os
import sys
from typing import List

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import CT_R, OxmlElement, parse_xml
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor, Inches, Pt

if sys.platform.startswith('win'):
    import pythoncom  # Required to initialize COM libraries for Windows
    import win32com.client  # Used for automating Microsoft Word

from .SharedEnums import ChangeStatusEnum
from . import DocGenInterface

from . import Logger

from docx.table import Table
from docx.text.paragraph import Paragraph

class DocxGen(DocGenInterface.DocGenInterface):

    def __init__(self, title: str, new_version: str,  time: str, origin_version: str = "",
                        origin_dds_types_files: list = None,
                        new_dds_types_files: list = None,
                        origin_dds_topics_types_mapping: str = "",
                        new_dds_topics_types_mapping: str = "",
                        is_compare_by_name: bool = False,
                        is_ignore_id: bool = False,
                        is_summary_report=False,
                        average_str_length: int = 1,
                        average_sequence_length: int = 1):
        """
        This class allows creating one of two documents; an ICD (created by docGen) and a comparison document (created
        by DocCompare). The difference is in the title page; the version in the ICD is added on its on, but in the
        comparison document both versions are displayed in a table along with the files/folders used by each version

        IMPORTANT:
        WHEN origin_version COMES EMPTY , WE KNOW IT'S AN ICD, and new_version IS USED AS THE VERSION.
        WHEN origin_version IS NOT EMPTY, WE KNOW IT'S A COMPARISON DOCUMENT.

        """
        super().__init__(title=title,
                         new_version=new_version,
                         time=time,
                         origin_version=origin_version,
                         origin_dds_types_files=origin_dds_types_files,
                         new_dds_types_files=new_dds_types_files,
                         origin_dds_topics_types_mapping=origin_dds_topics_types_mapping,
                         new_dds_topics_types_mapping=new_dds_topics_types_mapping)
        self.logger = Logger.add_logger(__name__)

        self.approx_width = Inches(len('Hierarchy') * 0.09)  # approximating each of 12pt font chars to around 0.09 Inch

        # Create a new '.docx' file based on the styling info from the template
        self.document = None
        self.doc_template_path = DocxGen.get_doc_template_path()
        if not self.doc_template_path:
            self.logger.fatal("Template.docx wasn't found. Cannot create doc")
            raise Exception("Template.docx wasn't found. Cannot create doc")
        else:
            self.logger.debug("doc template path at " + self.doc_template_path)
            self.document = docx.Document(self.doc_template_path)

        # Clear the document
        self.document._body.clear_content()
        self.add_doc_title_page(title=title,
                         new_version=new_version,
                         time=time,
                         origin_version=origin_version,
                         origin_dds_types_files=origin_dds_types_files,
                         new_dds_types_files=new_dds_types_files,
                         origin_dds_topics_types_mapping=origin_dds_topics_types_mapping,
                         new_dds_topics_types_mapping=new_dds_topics_types_mapping,
                         is_compare_by_name=is_compare_by_name,
                         is_ignore_id=is_ignore_id, is_summary_report=is_summary_report)

        if origin_version != "":
            # When this is a comparison document, add the color legend table
            self.add_color_legend_table()
            self.add_note_to_user()
        else:
            # When this is an ICD:
            self.add_volumes_note_to_user(average_str_length, average_sequence_length)

        self.add_toc_page()

    @staticmethod
    def get_doc_template_path() -> str:
        doc_name = "Template.docx"
        local_path = os.path.join(os.path.dirname(os.path.realpath(__file__)), doc_name)
        cwd_path = os.path.join(os.path.join(os.getcwd(), doc_name))
        if os.path.exists(local_path):
            return local_path
        elif os.path.exists(cwd_path):
            return cwd_path
        else:
            # https://importlib-resources.readthedocs.io/en/latest/using.html
            from importlib_resources import files
            whl_path = files('').joinpath(doc_name)
            if whl_path.is_file():
                return str(whl_path)
        return None

    @staticmethod
    def prepare_type_files_text_for_versions_table(type_files_list: List[str]):
        result = ""
        for type_file in type_files_list:
            result += type_file + "\n"

        if result.endswith("\n"):
            result = result[:-1]

        return result

    def add_version_info(self, origin_or_new_text: str, version: str, dds_types_files: list, dds_topic_file: str):

        paragraph = self.document.add_paragraph(f'{origin_or_new_text} - {version}', style='Version_style')

        paragraph = self.document.add_paragraph('DDS Type Files', style='List Paragraph')
        paragraph.paragraph_format.left_indent = Inches(0.5)

        # Add sub-bullets
        for sub_item in dds_types_files:
            sub_p = self.document.add_paragraph(f'{sub_item}', style='empty_bullet_2')
            sub_p.paragraph_format.left_indent = Inches(1)

        if bool(dds_topic_file):
            paragraph = self.document.add_paragraph(f'DDS Topic File: {dds_topic_file}', style='List Paragraph')
            paragraph.paragraph_format.left_indent = Inches(0.5)

    def add_color_legend_table(self):
        # Add a page break
        self.document.add_page_break()

        # Add a few blank lines to push the table down
        for _ in range(2):  # adjust this number as needed
            self.document.add_paragraph()

        # Add a table with 5 rows and 1 column with 'Table Grid' style for gridlines
        table = self.document.add_table(rows=5, cols=1, style='Table Grid')

        # Center the table horizontally
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set the cell text for each row
        texts = ['Color Legend', 'No background color means there is no change',
                 'GREEN background means the field was ADDED',
                 'YELLOW background means field was MODIFIED',
                 'RED background means field was DELETED']

        # Set cell width
        for row in table.rows:
            for cell in row.cells:
                cell.width = Inches(5.0)  # set the width

        # Set cell colors
        color_mapping = [None, ChangeStatusEnum.NO_CHANGE, ChangeStatusEnum.ADDED,
                         ChangeStatusEnum.CHANGED, ChangeStatusEnum.DELETED]

        for i in range(5):
            paragraph = table.cell(i, 0).paragraphs[0]
            run = paragraph.add_run(texts[i])
            run.font.size = Pt(14)  # set font size to 14
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # center text horizontally
            # Set background color if color_mapping is not None
            if color_mapping[i] is not None:
                color_xml = ChangeStatusEnum.get_background_color(color_mapping[i])
                if color_xml is not None:
                    paragraph._p.get_or_add_pPr().append(color_xml)

    def add_centered_paragraph(self, text, bold=False, underline=False, font_size=20):
        """Utility function to add a centered paragraph with specified formatting."""
        p = self.document.add_paragraph()
        p.alignment = 1  # Centered alignment
        run = p.add_run(text)
        run.bold = bold
        run.italic = True
        run.underline = underline
        run.font.size = Pt(font_size)
        return p
    def add_italic_paragraph(self, text, bold=False, underline=False, font_size=20):
        """Utility function to add a centered paragraph with specified formatting."""
        p = self.document.add_paragraph()
        # Indent the entire paragraph from the left by 1 inch
        p.paragraph_format.left_indent = Inches(0.5)

        run = p.add_run(text)
        run.bold = bold
        run.italic = True
        run.underline = underline
        run.font.size = Pt(font_size)
        return p

    def add_term_description_bullets(self, term, description):
        p = self.document.add_paragraph()
        p.style = 'List Paragraph'

        # Underline the key term and set font size
        run_term = p.add_run(term)
        run_term.underline = True
        run_term.font.size = Pt(16)
        run_term.italic = True

        # Add the rest of the description and set font size
        run_description = p.add_run(description)
        run_description.underline = False
        run_description.font.size = Pt(16)
        run_description.italic = True

    def add_note_to_user(self):
        # This function adds a note to the user informing him that this document only shows the differences between
        # versions, not the entire ICD

        # Add an empty space between the color legend and the note
        for _ in range(3):  # adjust this number as needed
            self.document.add_paragraph()
        # Add the desired text
        self.add_centered_paragraph('NOTE', bold=True, underline=True)
        self.add_centered_paragraph('This document lists only topics with changes')
        self.add_centered_paragraph('between the two selected data model versions.')
        self.add_centered_paragraph('Unchanged topics have been omitted for clarity.')
        self.add_centered_paragraph('To view all topics, please use the DocGen UI ')
        self.add_centered_paragraph('or call the generate_document API.')

    def add_volumes_note_to_user(self, average_str_length: int, average_sequence_length: int):
        self.document.add_page_break()
        # Add an empty space between the color legend and the note
        for _ in range(2):  # adjust this number as needed
            self.document.add_paragraph()
        self.add_centered_paragraph('Document Notes:', bold=True, underline=True)
        self.add_italic_paragraph('The document provides the following network volume metrics for each message.', font_size=16)

        self.add_term_description_bullets(
            "Minimum", ": The smallest byte size volume of the message.")
        self.add_term_description_bullets('Maximum',": The largest byte size volume of the message.")
        self.add_term_description_bullets("Average",
                                          f": a byte size volume of the message assuming strings are of size "
                                          f"{average_str_length} and sequence are of size {average_sequence_length}"
                                          f" (according to user-provided configuration).")

    def add_doc_title_page(self, title: str, new_version: str,  time: str, origin_version: str,
                        origin_dds_types_files: list = None,
                        new_dds_types_files: list = None,
                        origin_dds_topics_types_mapping: str = "",
                        new_dds_topics_types_mapping: str = "",
                        is_compare_by_name: bool = False,
                        is_ignore_id: bool = False,
                        is_summary_report=False):
        """
        This function adds a title page, with creation time, including header & footer,
        to be appended as a preamble section.
        This function  has two types of title page: one for an ICD (created by docGen) and one for a comparison document
        (created by DocCompare). The difference is that the version in the ICD is added on its own, but in the
        comparison document both versions are displayed in a table along with the files/folders used by each version

        IMPORTANT:
        WHEN origin_version COMES EMPTY , WE KNOW IT'S AN ICD, and new_version IS USED AS THE VERSION.
        WHEN origin_version IS NOT EMPTY, WE KNOW IT'S A COMPARISON DOCUMENT.

        """
        self.logger.debug(self.add_doc_title_page.__name__)

        # Add the title:
        self.document.add_heading(title, 0)
        if origin_version == "":
            # When this is an ICD - add the versions. (if it's not, the versions will be added below the creation time)
            par: Paragraph = self.document.add_paragraph('Version ' + new_version)
            par.style = self.document.styles['version']

        par: Paragraph = self.document.add_paragraph('Generated on ' + time)
        par.style = self.document.styles['Subtitle']

        if origin_version != "":
            # When this is A COMPARISON DOCUMENT, add report type and comparison method:
            self.add_new_page()
            report_type = "Complete"
            if is_summary_report:
                report_type = "Summary"
            # create an empty paragraph with no style
            par: Paragraph = self.document.add_paragraph()
            # add the first part of the text with the desired style
            par.add_run('Document Type', style='bold_underline_14')
            # add the rest of the text with no style
            run  = par.add_run(f': {report_type} Report')
            run.font.size = Pt(14)

            # Tell the user if we are comparing by id or by name (and if we're ignoring the ID)
            comparison_type = "ID"
            if is_compare_by_name:
                comparison_type = "Name"
            ignore_id_text = ""
            if is_ignore_id and is_compare_by_name:
                ignore_id_text = " (Ignoring IDs)"
            # create an empty paragraph with no style
            par: Paragraph = self.document.add_paragraph()
            # add the first part of the text with the desired style
            par.add_run('Comparison Type', style='bold_underline_14')
            # add the rest of the text with no style
            run  = par.add_run(f': {comparison_type} base{ignore_id_text}')
            run.font.size = Pt(14)

            self.add_version_info("Origin", origin_version, origin_dds_types_files, origin_dds_topics_types_mapping)
            self.add_version_info("Revised", new_version, new_dds_types_files, new_dds_topics_types_mapping)

    def add_toc_page(self):
        """
        This function adds a table of contents with hyperlink.
        """
        self.logger.debug(self.add_toc_page.__name__)
        self.add_new_page()
        # https://stackoverflow.com/questions/18595864/python-create-a-table-of-contents-with-python-docx-lxml
        par = self.document.add_paragraph("Table of Content", "Table of Content")
        run = par.add_run()
        fldChar = OxmlElement('w:fldChar')  # creates a new element
        fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
        instrText.text = 'TOC \\o "1-2" \\h \\z \\u'  # change 1-3 depending on heading levels you need

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:t')
        fldChar3.text = "Right-click to update field."
        fldChar2.append(fldChar3)

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        r_element = run._r
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar4)

    def add_chapter(self, section_title, parent_change_status=ChangeStatusEnum.NO_CHANGE, level=1):
        """
        This function adds a new chapter, on a new page.
        :param section_title: section name.
        :param parent_change_status: The status of changes made to the parent item.
        """
        self.logger.debug(self.add_chapter.__name__)

        # Since the only change in chapter level could be in the description, do not allow the heading get yellow bg
        if parent_change_status == ChangeStatusEnum.CHANGED or level == 2:
            parent_change_status = ChangeStatusEnum.NO_CHANGE

        # Add section heading with background color according to the change status
        color_xml = ChangeStatusEnum.get_background_color(parent_change_status)
        heading = self.document.add_heading(level=level)
        run = heading.add_run(section_title)

        if color_xml is not None:
            heading._p.get_or_add_pPr().append(color_xml)

    def add_type_to_chapter(self, dds_type, parent_change_status=ChangeStatusEnum.NO_CHANGE):
        # Since the only change in chapter level could be in the description, do not allow the heading get yellow bg
        if parent_change_status == ChangeStatusEnum.CHANGED:
            parent_change_status = ChangeStatusEnum.NO_CHANGE

        color_xml = ChangeStatusEnum.get_background_color(parent_change_status)
        sub_section_par = self.document.add_paragraph()
        sub_section_par.add_run('Type').underline = True
        sub_section_par.add_run(": " + dds_type)
        if color_xml is not None:
            sub_section_par._p.get_or_add_pPr().append(color_xml)

    @staticmethod
    def set_repeat_table_header(row):
        """ set repeat table row on every new page
        """
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), "true")
        trPr.append(tblHeader)
        return row

    def add_paragraph(self, paragraph_text):
        paragraph = self.document.add_paragraph(paragraph_text, style='Version_style')

    def add_table_header(self, listOfTitles, bLongHeader=True, color="grey"):
        """
        This function creates a table, and adds a table header with the titles.

        :param listOfTitles: the titles of the columns at the table header.
        :param bLongHeader: not in use, inherited from base class. May be needed in
                            other implementation environments. Defaults to True.
        :param color: the color of the table header, for better visualization.
                      Defaults to grey.
        """
        self.logger.debug(self.add_table_header.__name__)
        table: Table = self.document.add_table(rows=1, cols=len(listOfTitles), style='Table Grid')
        table.style = self.document.styles['Table Grid']
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        table_header_cells = table.rows[0].cells
        for i, title in enumerate(listOfTitles):
            table_header_cells[i].text = title
            table_header_cells[i].paragraphs[0].runs[0].font.bold = True
            if i == 0:  # If this is the first cell
                table_header_cells[i].width = self.approx_width  # Set the width
        DocxGen.set_repeat_table_header(table.rows[0])

        return table

    def add_table_row(self, theTable, cells_text, align='c', change_status=ChangeStatusEnum.NO_CHANGE):  # align=centered
        """
        This function adds a row for a table.
        :param theTable: the table object.
        :param cells_text: the data for the columns - as text.
        :param align:  alignment - centered (not in use here)
        :param bg_color: the background color for the row. Options are: ChangeStatusEnum.NO_CHANGE, ChangeStatusEnum.DELETED, ChangeStatusEnum.ADDED, ChangeStatusEnum.CHANGED, ChangeStatusEnum.TBD
        """
        self.logger.debug(self.add_table_row.__name__)
        row_cells = theTable.add_row().cells
        for i, cell_content in enumerate(cells_text):
            row_cells[i].text += cell_content
            color_xml = ChangeStatusEnum.get_background_color(change_status)
            if color_xml is not None:
                row_cells[i]._tc.get_or_add_tcPr().append(color_xml)
            if i == 0:  # If this is the first cell
                row_cells[i].width = self.approx_width  # Set the width

        if cells_text[0] == "":
            # When this element is on the basic level of the type, write the name in bold.
            row_cells[1].paragraphs[0].runs[0].font.bold = True

    def add_new_page(self):
        """
        This function adds a new page.
        """
        self.logger.debug(self.add_new_page.__name__)
        self.document.add_page_break()

    def add_section(self):
        """
        This function adds a new section. Not implemented.
        """
        self.logger.debug(self.add_section.__name__)
        pass

    def add_description(self, descr, parent_change_status=ChangeStatusEnum.NO_CHANGE):
        """
        This function adds a description (within a new section/subsection).
        :param descr: The description text to add.
        :param parent_change_status: The status of changes made to the parent item.
        """
        self.logger.debug(self.add_description.__name__)

        # Add description with background color according to the change status
        color_xml = ChangeStatusEnum.get_background_color(parent_change_status)
        description = self.document.add_paragraph()
        run = description.add_run('Description')
        run.underline = True
        run = description.add_run(": " + descr)
        if color_xml is not None:
            description._p.get_or_add_pPr().append(color_xml)

        self.add_new_line()

    def add_message_volumes(self, min_volume, avg_volume, max_volume):
        description = self.document.add_paragraph()
        run = description.add_run('Message Network Metrics')
        run.underline = True

        max_volume_str = f"{max_volume} Bytes"
        if max_volume == -1:
            max_volume_str = "Unbounded"
        volumes_str = f"   Minimum: {min_volume} Bytes.   Average: {avg_volume} Bytes.   Maximum: {max_volume_str}."
        run = description.add_run(": " + volumes_str)

    def add_new_line(self):
        """
        This function adds a new line.
        """
        self.logger.debug(self.add_new_line.__name__)
        # Add an empty line
        self.document.add_paragraph("")

    def generate_doc(self, output_file_name: str, temp_folder: str):
        """
        This function invokes the generation of the docx file then saves it in the requested folder.
        """
        self.logger.debug(self.generate_doc.__name__)
        output_file_name = f'{output_file_name}.docx'
        temp_output_path = os.path.abspath(os.path.join(temp_folder, output_file_name))
        self.logger.info(f"Generating {output_file_name}")

        try:
            self.document.save(temp_output_path)
            self.logger.debug(f"Initial writing to temporary folder succeeded.")
        except Exception as err:
            self.logger.error(f"The operation of saving into temporary folder has FAILED", exc_info=True)
            return

        # Update TOC
        try:
            self.update_toc(temp_output_path)
            self.logger.debug("File was saved at temporary folder")
        except Exception as err:
            self.logger.warn("Office is not installed. Cannot update table of content. User will have to perform "
                             "manually (F9 after opening document). Error = {}".format(str(err)))

    def finalize_doc(self):
        self.logger.debug(self.finalize_doc.__name__)
        pass
        # TODO set_autofit takes a long time; Could take minutes in large documents. Need a better solution
        # self.set_autofit()

    # Function to update the Table of Contents (TOC) in a Word document
    # This function works cross-platform (Windows and Linux/macOS)
    def update_toc(self, temp_path):
        # Check if the current platform is Windows
        if sys.platform.startswith('win'):

            pythoncom.CoInitialize()  # Initialize the COM library for the current thread
            word = None  # Placeholder for the Word application instance

            try:
                file_path = os.path.abspath(temp_path)  # Get the absolute path of the document
                # Start a new instance of Microsoft Word
                word = win32com.client.DispatchEx("Word.Application")
                word.DisplayAlerts = False  # Disable pop-up alerts from Word
                # Open the document
                doc = word.Documents.Open(file_path)
                # Update the first Table of Contents (if present)
                doc.TablesOfContents(1).Update()
                # Save and close the document
                doc.Close(SaveChanges=True)
            finally:
                # Ensure Word is properly closed even if an error occurs
                if word:
                    word.Quit()
                pythoncom.CoUninitialize()  # Uninitialize COM library to free resources
        # There's no 'else' here because docx is not suported in linux

    def set_autofit(self):
        """
        Make all table autofit to content
        """
        self.logger.debug(self.set_autofit.__name__)
        # TODO check if can set auto fit to window
        # https://github.com/python-openxml/python-docx/issues/209
        for t_idx, table in enumerate(self.document.tables):
            self.document.tables[t_idx].autofit = True
            self.document.tables[t_idx].allow_autofit = True
            self.document.tables[t_idx]._tblPr.xpath("./w:tblW")[0].attrib[
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"] = "auto"
            for row_idx, r_val in enumerate(self.document.tables[t_idx].rows):
                for cell_idx, c_val in enumerate(self.document.tables[t_idx].rows[row_idx].cells):
                    self.document.tables[t_idx].rows[row_idx].cells[cell_idx]._tc.tcPr.tcW.type = 'auto'
                    self.document.tables[t_idx].rows[row_idx].cells[cell_idx]._tc.tcPr.tcW.w = 0

# https://github.com/python-openxml/python-docx/commit/65db85311e9de6e50add607be169e57f8fcc7591
#         cp = self.document.paragraphs[0]
#         print(cp.text)
#         new_paragraph = cp.insert_paragraph_before('barfoo')


# TODO change table styling
