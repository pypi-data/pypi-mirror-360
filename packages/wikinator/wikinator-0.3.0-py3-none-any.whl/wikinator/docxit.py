# This file is original taken from https://github.com/haesleinhuepf/docx2markdown
# - src/docx2markdown/_docx_to_markdown.py
# It is being altered here to provide results needed for this project.
# original license: https://github.com/haesleinhuepf/docx2markdown/blob/main/LICENSE
# included here for completeness
#---
# BSD 3-Clause License
# Copyright (c) 2024, Robert Haase, ScaDS.AI, Uni Leipzig
# Redistribution and use in source and binary forms, with or without
# modification, are permitted provided that the following conditions are met:
# 1. Redistributions of source code must retain the above copyright notice, this
#    list of conditions and the following disclaimer.
# 2. Redistributions in binary form must reproduce the above copyright notice,
#    this list of conditions and the following disclaimer in the documentation
#    and/or other materials provided with the distribution.
# 3. Neither the name of the copyright holder nor the names of its
#    contributors may be used to endorse or promote products derived from
#    this software without specific prior written permission.
# THIS SOFTWARE IS PROVIDED BY THE COPYRIGHT HOLDERS AND CONTRIBUTORS "AS IS"
# AND ANY EXPRESS OR IMPLIED WARRANTIES, INCLUDING, BUT NOT LIMITED TO, THE
# IMPLIED WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE ARE
# DISCLAIMED. IN NO EVENT SHALL THE COPYRIGHT HOLDER OR CONTRIBUTORS BE LIABLE
# FOR ANY DIRECT, INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY, OR CONSEQUENTIAL
# DAMAGES (INCLUDING, BUT NOT LIMITED TO, PROCUREMENT OF SUBSTITUTE GOODS OR
# SERVICES; LOSS OF USE, DATA, OR PROFITS; OR BUSINESS INTERRUPTION) HOWEVER
# CAUSED AND ON ANY THEORY OF LIABILITY, WHETHER IN CONTRACT, STRICT LIABILITY,
# OR TORT (INCLUDING NEGLIGENCE OR OTHERWISE) ARISING IN ANY WAY OUT OF THE USE
# OF THIS SOFTWARE, EVEN IF ADVISED OF THE POSSIBILITY OF SUCH DAMAGE.
#---
# All changes in the this version are Copyright (c) 2025, Paul Philion, Acme Rocket Company
# under the provided MIT license.

import docx
import os
from lxml import etree
from pathlib import Path

def docx_to_markdown(docx_file, output_md):
    """Convert a .docx file to a Markdown file and a subfolder of images."""

    folder = str(Path(output_md).parent)
    image_folder = str(Path(output_md).parent / "images")

    doc = docx.Document(docx_file)

    paragraphs = list(doc.paragraphs)
    tables = list(doc.tables)
    markdown = []

    # save all images
    images = {}
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_filename = save_image(rel.target_part, image_folder)
            images[rel.rId] = image_filename[len(folder)+1:]
            #print("image file", image_filename, "-->", images[rel.rId])

    #print("images", images)

    for block in doc.element.body:
        if block.tag.endswith('p'):  # Handle paragraphs
            paragraph = paragraphs.pop(0)  # Match current paragraph
            md_paragraph = ""

            ### switching on paragraph.style.name

            style_name = paragraph.style.name

            if "List" in style_name:
                prefix = get_bullet_point_prefix(paragraph)
                md_paragraph = prefix  # Markdown syntax for bullet points
            elif "Heading 1" in style_name:
                md_paragraph = "# "
            elif "Heading 2" in style_name:
                md_paragraph = "## "
            elif "Heading 3" in style_name:
                md_paragraph = "### "
            elif "Heading 4" in style_name:
                md_paragraph = "#### "
            elif "Heading 5" in style_name:
                md_paragraph = "##### "

            elif "Normal" or "normal" in style_name:
                md_paragraph = ""
            else:
                print("Unsupported style:", style_name)

            md_paragraph += parse_run(paragraph, images)

            markdown.append(md_paragraph)

        elif block.tag.endswith('tbl'):  # Handle tables (if present)
            table = tables.pop(0)  # Match current table
            table_text = ""
            for i, row in enumerate(table.rows):
                table_text += "| " + " | ".join(cell.text.strip() for cell in row.cells) + " |\n"
                if i == 0:
                    table_text += "| " + " | ".join("---" for _ in row.cells) + " |\n"

            markdown.append(table_text)

        elif block.tag.endswith('sectPr') or block.tag.endswith('sdt'):
            # ignore
            pass
        else:
            print("Unsupported block:", docx_file, block.tag)

    # Write to Markdown file
    with open(output_md, "w", encoding="utf-8") as md_file:
        md_file.write("\n\n".join(markdown))


def extract_r_embed(xml_string):
    """
    Extract the value of r:embed from the given XML string.

    :param xml_string: The XML content as a string.
    :return: The value of r:embed or None if not found.
    """
    # Parse the XML
    root = etree.fromstring(xml_string)

    # Define the namespaces
    namespaces = {
        'a': "http://schemas.openxmlformats.org/drawingml/2006/main",
        'r': "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        'pic': "http://schemas.openxmlformats.org/drawingml/2006/picture",
    }

    # Use XPath to find the <a:blip> element with r:embed
    blip = root.find(".//a:blip", namespaces=namespaces)

    # Extract the r:embed attribute value
    if blip is not None:
        return blip.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
    return None

def save_image(image_part, output_folder):
    """Save an image to the output folder and return the filename."""
    os.makedirs(output_folder, exist_ok=True)
    image_filename = os.path.join(output_folder, os.path.basename(image_part.partname))
    with open(image_filename, "wb") as img_file:
        img_file.write(image_part.blob)
    return str(image_filename).replace("\\", "/")


def get_list_level(paragraph):
    """Determine the level of a bullet point or numbered list item."""
    # Access the raw XML of the paragraph
    p = paragraph._element
    numPr = p.find(".//w:numPr", namespaces=p.nsmap)
    if numPr is not None:
        ilvl = numPr.find(".//w:ilvl", namespaces=p.nsmap)
        if ilvl is not None:
            return int(ilvl.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"))
    return 0

def get_bullet_point_prefix(paragraph):
    """Determine the Markdown prefix for a bullet point based on its indentation level."""
    level = get_list_level(paragraph)
    return "  " * level + "- "  # Use Markdown syntax for nested lists

def parse_run(run, images):
    """Go through document objects recursively and return markdown."""
    sub_parts = list(run.iter_inner_content())
    text = ""
    for s in sub_parts:
        if isinstance(s, str):
            text += s
        elif isinstance(s, docx.text.run.Run):
            text += parse_run(s, images)
        elif isinstance(s, docx.text.hyperlink.Hyperlink):
            text += f"[{s.text}]({s.address})"
        elif isinstance(s, docx.drawing.Drawing):
            rId = extract_r_embed(s._element.xml)
            image_url = images[rId]
            text += f"![]({image_url})"
        else:
            print("unknown run type", s)

    if isinstance(run, docx.text.run.Run):
        if run.bold:
            text = f"**{text}**"
        if run.italic:
            text = f"*{text}*"
        if run.underline:
            text = f"__{text}__"
        if run.font.strike:
            text = f"~~{text}~~"
        # check .font for monospacing
        # check style
        if run.font.name == "Courier New": # more fonts!
            text = f"`{text}`"

    return text
