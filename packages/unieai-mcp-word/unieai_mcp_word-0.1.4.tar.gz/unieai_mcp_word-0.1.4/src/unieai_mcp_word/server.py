from fastmcp import FastMCP
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import tempfile
from datetime import datetime
import os, shutil, re, uuid

PATTERN = re.compile(r"{{\s*(\w+)\s*}}")  # 捕捉 {{ key }} 或 {{key}}

def main():
    mcp = FastMCP("unieai-mcp-word-stdio")

    @mcp.tool()
    def write_data_to_word_with_custom(data: dict) -> str:
        """
        UnieAI 專用客戶回饋報告 Word 模板
        將 data 的 key/value 套入 Word 範本中的 {{key}} 位置，
        完成後回傳下載超連結（Markdown 內嵌）。
        """
        outputpath = (
            "/app/data/storage/unieai-mcp-word/customer_feedback_report_"
            + datetime.now().strftime("%Y%m%d")
            + "_"
            + str(uuid.uuid4())
            + ".docx"
        )
        params = {
            "filepath": "/app/data/storage/unieai-mcp-word/customer_feedback_report_temp.docx",
            "outputpath": outputpath,
            "data_map": data,
            "desc": "客戶回饋報告 Word檔案下載"
        }
        return fill_word_with_context(params)



    mcp.run(transport="stdio")



def fill_word_with_context(params: dict) -> str:
    src, dst, ctx = params["filepath"], params["outputpath"], params["data_map"]
    os.makedirs(os.path.dirname(dst), exist_ok=True)
    shutil.copy(src, dst)                             # 複製範本 → 新檔:contentReference[oaicite:5]{index=5}
    doc = Document(dst)

    def _replace_in_runs(runs):
        for run in runs:
            if "{{" not in run.text:
                continue
            new_text = PATTERN.sub(lambda m: str(ctx.get(m.group(1), m.group(0))), run.text)
            run.text = new_text

    # ── 1) 逐段落
    for p in doc.paragraphs:
        _replace_in_runs(p.runs)

    # ── 2) 逐表格儲存格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for cp in cell.paragraphs:
                    _replace_in_runs(cp.runs)         # 表格中文字替換:contentReference[oaicite:6]{index=6}

    # ── 3) 頁眉／頁腳
    for sect in doc.sections:
        for p in sect.header.paragraphs:
            _replace_in_runs(p.runs)                  # 讀取 header API:contentReference[oaicite:7]{index=7}
        for p in sect.footer.paragraphs:
            _replace_in_runs(p.runs)

    doc.save(dst)
    link = "https://office-mcp-dl.unieai.com/" + os.path.basename(dst)
    #link = dst
    return f"*[{params.get('desc', 'Word 檔案下載')}]({link})*"







if __name__ == "__main__":
    main()
