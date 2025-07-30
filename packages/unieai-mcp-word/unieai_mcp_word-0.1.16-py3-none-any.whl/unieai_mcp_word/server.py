from fastmcp import FastMCP
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import tempfile
from datetime import datetime
import os, shutil, re, uuid

#PATTERN = re.compile(r"{\s*(\w+)\s*}")  # 捕捉 {{ key }} 或 {{key}}
PATTERN = re.compile(r'\{(\w+)\}')

def main():
    mcp = FastMCP("unieai-mcp-word-stdio")

    @mcp.tool()
    def write_data_to_word_with_custom(data: dict) -> str:
        """
        UnieAI 專用客戶回饋報告 Word 模板
        將 data 的 key/value 套入 Word 範本中的 {{key}} 位置，
        完成後回傳下載超連結（Markdown 內嵌）。
        """
        outputpath = (
            #"/app/data/storage/unieai-mcp-word/customer_feedback_report_"
            "D:/customer_feedback_report_"
            + datetime.now().strftime("%Y%m%d")
            + "_"
            + str(uuid.uuid4())
            + ".docx"
        )
        params = {
            #"filepath": "/app/data/storage/unieai-mcp-word/customer_feedback_report_temp.docx",
            "filepath": "D:/customer_feedback_report_temp.docx",
            "outputpath": outputpath,
            "data_map": data,
            "desc": "客戶回饋報告 Word檔案下載"
        }
        return fill_word_with_context(params)



    mcp.run(transport="stdio")



def fill_word_with_context(params: dict) -> str:
    src, dst, ctx = params["filepath"], params["outputpath"], params["data_map"]
    os.makedirs(os.path.dirname(dst), exist_ok=True)
    shutil.copy(src, dst)                             # 複製範本 → 新檔:contentReference[oaicite:5]{index=5}
    doc = Document(dst)

    test_log = "test_log : "

    # def _replace_in_runs(runs):
    #     nonlocal test_log
    #     for run in runs:
    #         print("run.text:", run.text)
    #         test_log = test_log + "run.text:" + run.text + "\n"
    #         if "{" not in run.text:
    #             continue
    #         new_text = PATTERN.sub(lambda m: str(ctx.get(m.group(1), m.group(0))), run.text)
    #         print("new_text:", new_text)
    #         test_log = test_log + "new_text:" + new_text + "\n"
    #         print("--------------------------------")
    #         run.text = new_text


    def _replace_in_runs(paragraph):
    """
    替換段落中所有形如 {key} 的佔位符為 ctx 中對應的值，
    並保留原有的樣式設定。
    """
    # 合併所有 run 的文字為一個字串
    full_text = ''.join(run.text for run in paragraph.runs)
    
    # 找出所有的 {key} 佔位符
    matches = PATTERN.finditer(full_text)
    
    # 反向遍歷 matches，以便在替換時不影響尚未處理的部分
    for match in reversed(list(matches)):
        key = match.group(1)
        value = ctx.get(key, match.group(0))  # 若 ctx 中沒有對應的 key，則保留原始佔位符
        
        # 計算佔位符在 full_text 中的位置
        start_pos, end_pos = match.span()
        
        # 找出對應的 run
        start_run_index = sum(len(run.text) for run in paragraph.runs[:start_pos])
        end_run_index = sum(len(run.text) for run in paragraph.runs[:end_pos])
        
        # 替換對應位置的文字
        paragraph.runs[start_run_index].text = paragraph.runs[start_run_index].text[:start_pos] + value + paragraph.runs[start_run_index].text[end_pos:]

    # ── 1) 逐段落
    for p in doc.paragraphs:
        _replace_in_runs(p)

    # ── 2) 逐表格儲存格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for cp in cell.paragraphs:
                    _replace_in_runs(cp)         # 表格中文字替換:contentReference[oaicite:6]{index=6}

    # ── 3) 頁眉／頁腳
    for sect in doc.sections:
        for p in sect.header.paragraphs:
            _replace_in_runs(p)                  # 讀取 header API:contentReference[oaicite:7]{index=7}
        for p in sect.footer.paragraphs:
            _replace_in_runs(p)

    doc.save(dst)
    #link = "https://office-mcp-dl.unieai.com/unieai-mcp-word/" + os.path.basename(dst)
    link = dst
    #response = "*[報價單A檔案下載](" + link_url + ")*"
    return f"*[{params.get('desc', 'Word 檔案下載')}]({link})* \n\n{test_log}"







if __name__ == "__main__":
    main()
