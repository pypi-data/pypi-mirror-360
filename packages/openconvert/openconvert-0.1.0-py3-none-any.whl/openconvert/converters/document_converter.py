"""
Document converter module for handling document format conversions.
"""

import logging
import os
import tempfile
from pathlib import Path
from typing import Union, Optional, Dict, Any

logger = logging.getLogger(__name__)

# Define supported conversions
SUPPORTED_CONVERSIONS = {
    'txt': ['pdf', 'docx', 'rtf', 'md', 'html', 'csv'],
    'docx': ['pdf', 'txt', 'rtf', 'html', 'md', 'epub'],
    'pdf': ['docx', 'txt', 'jpg', 'png', 'epub', 'html'],
    'html': ['pdf', 'docx', 'txt', 'md'],
    'md': ['pdf', 'docx', 'html', 'txt'],
    'rtf': ['docx', 'pdf', 'txt'],
    'csv': ['xlsx', 'json', 'txt', 'xml', 'sql'],
    'xlsx': ['csv', 'json', 'xml', 'sql'],
    'epub': ['pdf', 'docx', 'txt', 'html']
}

def convert(
    filepath: Union[str, Path],
    source_format: str,
    target_format: str,
    output_path: Union[str, Path],
    options: Optional[Dict[str, Any]] = None
) -> str:
    """
    Convert a document from one format to another.
    
    Args:
        filepath: Path to the source document file
        source_format: Source document format
        target_format: Target document format
        output_path: Path to save the converted document
        options: Additional conversion options
        
    Returns:
        Path to the converted document
        
    Raises:
        ValueError: If the conversion is not supported
        RuntimeError: If the conversion fails
    """
    if options is None:
        options = {}
    
    # Check if conversion is supported
    if target_format not in SUPPORTED_CONVERSIONS.get(source_format, []):
        raise ValueError(f"Conversion from {source_format} to {target_format} is not supported")
    
    filepath = Path(filepath)
    output_path = Path(output_path)
    
    # Route to appropriate conversion method
    try:
        if source_format == 'txt':
            return _convert_from_txt(filepath, target_format, output_path, options)
        elif source_format == 'docx':
            return _convert_from_docx(filepath, target_format, output_path, options)
        elif source_format == 'pdf':
            return _convert_from_pdf(filepath, target_format, output_path, options)
        elif source_format == 'html':
            return _convert_from_html(filepath, target_format, output_path, options)
        elif source_format == 'md':
            return _convert_from_md(filepath, target_format, output_path, options)
        elif source_format == 'rtf':
            return _convert_from_rtf(filepath, target_format, output_path, options)
        elif source_format == 'csv':
            return _convert_from_csv(filepath, target_format, output_path, options)
        elif source_format == 'xlsx':
            return _convert_from_xlsx(filepath, target_format, output_path, options)
        elif source_format == 'epub':
            return _convert_from_epub(filepath, target_format, output_path, options)
        else:
            raise ValueError(f"Unsupported source format: {source_format}")
    
    except Exception as e:
        logger.error(f"Error converting {filepath} to {target_format}: {str(e)}")
        raise RuntimeError(f"Failed to convert {filepath} to {target_format}: {str(e)}")

def _convert_from_txt(
    filepath: Path,
    target_format: str,
    output_path: Path,
    options: Dict[str, Any]
) -> str:
    """Convert from TXT to other formats."""
    if target_format == 'pdf':
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            
            # Read the text file
            with open(filepath, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            # Create PDF
            doc = SimpleDocTemplate(str(output_path), pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Split text into paragraphs
            paragraphs = text_content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    p = Paragraph(para.replace('\n', '<br/>'), styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 12))
            
            doc.build(story)
            
        except ImportError:
            # Alternative method using pandoc
            _convert_using_pandoc(filepath, target_format, output_path)
    
    elif target_format == 'docx':
        try:
            from docx import Document
            
            # Read the text file
            with open(filepath, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            # Create DOCX
            doc = Document()
            
            # Split text into paragraphs
            paragraphs = text_content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para)
            
            doc.save(str(output_path))
            
        except ImportError:
            # Alternative method using pandoc
            _convert_using_pandoc(filepath, target_format, output_path)
    
    elif target_format == 'rtf':
        # Direct conversion to RTF without pandoc
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            # Create a simple RTF file
            rtf_header = r"{\rtf1\ansi\ansicpg1252\cocoartf2580\cocoasubrtf220"
            rtf_header += r"{\fonttbl\f0\fswiss\fcharset0 Helvetica;}"
            rtf_header += r"{\colortbl;\red255\green255\blue255;}"
            rtf_header += r"{\*\expandedcolortbl;;}"
            rtf_header += r"\margl1440\margr1440\vieww11520\viewh8400\viewkind0"
            rtf_header += r"\pard\tx720\tx1440\tx2160\tx2880\tx3600\tx4320\tx5040\tx5760\tx6480\tx7200\tx7920\tx8640\pardirnatural\partightenfactor0"
            rtf_header += r"\f0\fs24 \cf0 "
            
            # Prepare content for RTF
            rtf_content = text_content.replace("\n", "\\par\n")
            rtf_content = rtf_content.replace("\\", "\\\\")
            rtf_content = rtf_content.replace("{", "\\{")
            rtf_content = rtf_content.replace("}", "\\}")
            
            # Write the RTF file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(rtf_header + rtf_content + "}")
            
        except Exception as e:
            logger.warning(f"Direct RTF conversion failed: {str(e)}. Trying with pandoc...")
            # Fallback to pandoc
            try:
                _convert_using_pandoc(filepath, target_format, output_path)
            except Exception:
                # Create a very simple RTF as a last resort
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(r'{\rtf1\ansi\deff0{\fonttbl{\f0 Times New Roman;}}{\colortbl;\red0\green0\blue0;}\f0\fs24\cf1 ')
                    with open(filepath, 'r', encoding='utf-8') as in_f:
                        for line in in_f:
                            f.write(line.replace('\n', '\\par\n'))
                    f.write('}')
    
    elif target_format in ['html', 'md']:
        # Use pandoc for these conversions
        _convert_using_pandoc(filepath, target_format, output_path)
    
    elif target_format == 'csv':
        # Simple conversion - each line becomes a row with a single column
        with open(filepath, 'r', encoding='utf-8') as f:
            text_content = f.read()
        
        with open(output_path, 'w', encoding='utf-8', newline='') as f:
            import csv
            writer = csv.writer(f)
            for line in text_content.split('\n'):
                writer.writerow([line])
    
    return str(output_path)

def _convert_from_docx(
    filepath: Path,
    target_format: str,
    output_path: Path,
    options: Dict[str, Any]
) -> str:
    """Convert from DOCX to other formats."""
    if target_format == 'txt':
        try:
            from docx import Document
            
            doc = Document(filepath)
            with open(output_path, 'w', encoding='utf-8') as f:
                for para in doc.paragraphs:
                    f.write(para.text + '\n')
                    if para.text.strip() == '':
                        f.write('\n')  # Add extra newline for paragraph breaks
        
        except ImportError:
            # Alternative method using pandoc
            _convert_using_pandoc(filepath, target_format, output_path)
    
    elif target_format in ['pdf', 'html', 'md', 'rtf', 'epub']:
        # Use pandoc for these conversions
        _convert_using_pandoc(filepath, target_format, output_path)
    
    return str(output_path)

def _convert_from_pdf(
    filepath: Path,
    target_format: str,
    output_path: Path,
    options: Dict[str, Any]
) -> str:
    """Convert from PDF to other formats."""
    if target_format == 'txt':
        try:
            import PyPDF2
            
            with open(filepath, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text = ''
                for page_num in range(len(pdf_reader.pages)):
                    text += pdf_reader.pages[page_num].extract_text() + '\n\n'
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
        
        except ImportError:
            # Alternative method using pdfminer
            try:
                from pdfminer.high_level import extract_text
                
                text = extract_text(str(filepath))
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text)
            
            except ImportError:
                # Last resort: try pandoc
                _convert_using_pandoc(filepath, target_format, output_path)
    
    elif target_format in ['docx', 'html', 'epub']:
        # Try using pandoc
        _convert_using_pandoc(filepath, target_format, output_path)
    
    elif target_format in ['jpg', 'png']:
        # Convert PDF to image
        try:
            from pdf2image import convert_from_path
            
            images = convert_from_path(filepath)
            
            # If there's only one page, save it directly
            if len(images) == 1:
                images[0].save(str(output_path))
            else:
                # If there are multiple pages, save them with sequential numbering
                base_path = output_path.with_suffix('')
                for i, image in enumerate(images):
                    page_path = f"{base_path}_page_{i+1}{output_path.suffix}"
                    image.save(page_path)
                
                # Return the directory containing all pages
                return str(base_path.parent)
        
        except ImportError:
            raise RuntimeError("pdf2image library is required for PDF to image conversion. Please install it.")
    
    return str(output_path)

def _convert_from_html(
    filepath: Path,
    target_format: str,
    output_path: Path,
    options: Dict[str, Any]
) -> str:
    """Convert from HTML to other formats."""
    if target_format == 'txt':
        try:
            from bs4 import BeautifulSoup
            
            with open(filepath, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f, 'html.parser')
                text = soup.get_text(separator='\n\n')
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
        
        except ImportError:
            # Alternative method using pandoc
            _convert_using_pandoc(filepath, target_format, output_path)
    
    elif target_format == 'pdf':
        # Try multiple methods for HTML to PDF conversion
        success = False
        
        # Method 1: Try using pdfkit (wkhtmltopdf)
        try:
            import pdfkit
            
            # Configure pdfkit
            options_dict = {
                'encoding': 'UTF-8',
                'quiet': ''
            }
            
            # Add any user-provided options
            if options:
                for key, value in options.items():
                    if key not in ['quality', 'resize']:  # Skip general options
                        options_dict[key] = value
            
            pdfkit.from_file(str(filepath), str(output_path), options=options_dict)
            success = True
            logger.info("Converted HTML to PDF using pdfkit")
        
        except (ImportError, Exception) as e:
            logger.warning(f"pdfkit conversion failed: {str(e)}. Trying next method...")
        
        # Method 2: Try using weasyprint
        if not success:
            try:
                from weasyprint import HTML
                
                HTML(filename=str(filepath)).write_pdf(str(output_path))
                success = True
                logger.info("Converted HTML to PDF using weasyprint")
            
            except (ImportError, Exception) as e:
                logger.warning(f"weasyprint conversion failed: {str(e)}. Trying next method...")
        
        # Method 3: Try using pandoc as a last resort
        if not success:
            try:
                _convert_using_pandoc(filepath, target_format, output_path)
                logger.info("Converted HTML to PDF using pandoc")
                success = True
            except Exception as e:
                logger.error(f"All HTML to PDF conversion methods failed: {str(e)}")
                raise RuntimeError(f"Failed to convert HTML to PDF. Please install pdfkit, weasyprint, or pandoc with LaTeX.")
    
    elif target_format == 'docx' or target_format == 'md':
        # Use pandoc for these conversions
        _convert_using_pandoc(filepath, target_format, output_path)
    
    return str(output_path)

def _convert_from_md(
    filepath: Path,
    target_format: str,
    output_path: Path,
    options: Dict[str, Any]
) -> str:
    """Convert from Markdown to other formats."""
    if target_format == 'txt':
        # Simple conversion - just strip markdown syntax
        with open(filepath, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Very basic markdown stripping
        import re
        
        # Remove headers
        text = re.sub(r'^#+\s+', '', md_content, flags=re.MULTILINE)
        
        # Remove bold/italic
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        
        # Remove links
        text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
        
        # Remove code blocks
        text = re.sub(r'```.*?```', '', text, flags=re.DOTALL)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
    
    elif target_format in ['pdf', 'docx', 'html']:
        # Use pandoc for these conversions
        _convert_using_pandoc(filepath, target_format, output_path)
    
    return str(output_path)

def _convert_from_rtf(
    filepath: Path,
    target_format: str,
    output_path: Path,
    options: Dict[str, Any]
) -> str:
    """Convert from RTF to other formats."""
    # Use pandoc for all RTF conversions
    _convert_using_pandoc(filepath, target_format, output_path)
    return str(output_path)

def _convert_from_csv(
    filepath: Path,
    target_format: str,
    output_path: Path,
    options: Dict[str, Any]
) -> str:
    """Convert from CSV to other formats."""
    if target_format == 'xlsx':
        try:
            import pandas as pd
            
            df = pd.read_csv(filepath)
            df.to_excel(output_path, index=False)
        
        except ImportError:
            raise RuntimeError("pandas library is required for CSV to XLSX conversion. Please install it.")
    
    elif target_format == 'json':
        try:
            import pandas as pd
            
            df = pd.read_csv(filepath)
            df.to_json(output_path, orient='records', indent=4)
        
        except ImportError:
            # Alternative method using csv and json modules
            import csv
            import json
            
            with open(filepath, 'r', encoding='utf-8', newline='') as f:
                reader = csv.DictReader(f)
                data = list(reader)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=4)
    
    elif target_format == 'txt':
        # Simple conversion - just copy the CSV content
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
    
    elif target_format == 'xml':
        try:
            import pandas as pd
            import dicttoxml
            
            df = pd.read_csv(filepath)
            data = df.to_dict(orient='records')
            xml = dicttoxml.dicttoxml(data, custom_root='data', attr_type=False)
            
            with open(output_path, 'wb') as f:
                f.write(xml)
        
        except ImportError:
            raise RuntimeError("pandas and dicttoxml libraries are required for CSV to XML conversion. Please install them.")
    
    elif target_format == 'sql':
        try:
            import pandas as pd
            import sqlite3
            
            df = pd.read_csv(filepath)
            
            # Create a temporary SQLite database
            conn = sqlite3.connect(str(output_path))
            
            # Get table name from file name or use default
            table_name = options.get('table_name', filepath.stem)
            
            # Write to SQL
            df.to_sql(table_name, conn, index=False)
            conn.close()
        
        except ImportError:
            raise RuntimeError("pandas library is required for CSV to SQL conversion. Please install it.")
    
    return str(output_path)

def _convert_from_xlsx(
    filepath: Path,
    target_format: str,
    output_path: Path,
    options: Dict[str, Any]
) -> str:
    """Convert from XLSX to other formats."""
    try:
        import pandas as pd
        
        # Read the Excel file
        df = pd.read_excel(filepath, sheet_name=options.get('sheet_name', 0))
        
        if target_format == 'csv':
            df.to_csv(output_path, index=False)
        
        elif target_format == 'json':
            df.to_json(output_path, orient='records', indent=4)
        
        elif target_format == 'xml':
            import dicttoxml
            
            data = df.to_dict(orient='records')
            xml = dicttoxml.dicttoxml(data, custom_root='data', attr_type=False)
            
            with open(output_path, 'wb') as f:
                f.write(xml)
        
        elif target_format == 'sql':
            import sqlite3
            
            # Create a temporary SQLite database
            conn = sqlite3.connect(str(output_path))
            
            # Get table name from file name or use default
            table_name = options.get('table_name', filepath.stem)
            
            # Write to SQL
            df.to_sql(table_name, conn, index=False)
            conn.close()
    
    except ImportError:
        raise RuntimeError("pandas library is required for XLSX conversions. Please install it.")
    
    return str(output_path)

def _convert_from_epub(
    filepath: Path,
    target_format: str,
    output_path: Path,
    options: Dict[str, Any]
) -> str:
    """Convert from EPUB to other formats."""
    # Use pandoc for all EPUB conversions
    _convert_using_pandoc(filepath, target_format, output_path)
    return str(output_path)

def _convert_using_pandoc(
    filepath: Path,
    target_format: str,
    output_path: Path
) -> None:
    """Use pandoc for document conversion."""
    try:
        import subprocess
        
        # Map our format names to pandoc format names
        format_map = {
            'txt': 'markdown',  # Use markdown as the input format for .txt files
            'md': 'markdown',
            'html': 'html',
            'docx': 'docx',
            'pdf': 'pdf',
            'rtf': 'rtf',
            'epub': 'epub'
        }
        
        source_ext = filepath.suffix.lower().lstrip('.')
        if source_ext == 'jpeg':
            source_ext = 'jpg'
        
        pandoc_source = format_map.get(source_ext, source_ext)
        pandoc_target = format_map.get(target_format, target_format)
        
        cmd = [
            'pandoc',
            '-f', pandoc_source,
            '-t', pandoc_target,
            '-o', str(output_path),
            str(filepath)
        ]
        
        # For debug purposes, log the command
        logger.debug(f"Running pandoc command: {' '.join(cmd)}")
        
        subprocess.run(cmd, check=True)
    
    except (ImportError, subprocess.SubprocessError) as e:
        raise RuntimeError(f"Failed to convert using pandoc: {str(e)}. Please install pandoc.") 