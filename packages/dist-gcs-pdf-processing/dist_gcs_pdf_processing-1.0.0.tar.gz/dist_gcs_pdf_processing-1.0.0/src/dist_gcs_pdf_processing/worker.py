"""
Worker: Concurrent file and page processing with per-page retries and global Gemini API throttling.
- File-level concurrency: MAX_CONCURRENT_FILES
- Per-page concurrency: PAGE_MAX_WORKERS
- Per-page retries: MAX_RETRIES
- Global Gemini API throttling: GEMINI_GLOBAL_CONCURRENCY
- All logs/prints/errors include per-file trace_id
"""
import os
from dist_gcs_pdf_processing.env import load_env_and_credentials

load_env_and_credentials()
os.environ["G_MESSAGES_DEBUG"] = "none"
os.environ["G_DEBUG"] = "fatal-warnings"
os.environ["PYTHONWARNINGS"] = "ignore"
import time
import tempfile
import shutil
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed, wait, FIRST_COMPLETED
from collections import namedtuple
from typing import List
# import redis 
# from prometheus_client import Counter, Histogram 
from .gcs_utils import download_from_gcs, upload_to_gcs, list_new_files
from .ocr import gemini_ocr_page
from .config import POLL_INTERVAL, DOC_BATCH_SIZE, PAGE_MAX_WORKERS, STAGING_DIR, PROCESSED_DIR
from pypdf import PdfReader, PdfWriter
import markdown2
from weasyprint import HTML
from docx import Document
# import openai 
import mimetypes
from logging.handlers import TimedRotatingFileHandler
import json
from datetime import datetime, timedelta
import uuid
import requests
import threading
import base64
import signal
import atexit
from .shared import GCS_LIMITER, GEMINI_LIMITER, RateLimiter
from queue import Queue, Empty
import concurrent.futures

# FILES_PROCESSED = Counter("files_processed_total", "Total files processed")
# PAGES_PROCESSED = Counter("pages_processed_total", "Total pages processed")
# PROCESSING_TIME = Histogram("file_processing_seconds", "Time spent processing files")

# redis_client = redis.Redis.from_url(REDIS_URL) if REDIS_URL else None

# Set up a logs directory and file handler for local logging
LOGS_DIR = os.path.join(os.path.dirname(__file__), "logs")
JSON_LOGS_DIR = os.path.join(LOGS_DIR, "json")
DEAD_LETTER_DIR = os.path.join(LOGS_DIR, "dead_letter")
os.makedirs(LOGS_DIR, exist_ok=True)
os.makedirs(JSON_LOGS_DIR, exist_ok=True)
os.makedirs(DEAD_LETTER_DIR, exist_ok=True)

# Set up daily rotating log file
log_file_path = os.path.join(LOGS_DIR, "worker.log")
file_handler = TimedRotatingFileHandler(log_file_path, when="midnight", backupCount=200)
file_handler.setFormatter(logging.Formatter('%(asctime)s %(levelname)s %(message)s'))
stream_handler = logging.StreamHandler()
stream_handler.setFormatter(logging.Formatter('%(asctime)s %(levelname)s %(message)s'))
logging.basicConfig(level=logging.INFO, handlers=[file_handler, stream_handler])
logger = logging.getLogger("dcpr.worker")

# JSON logging function
def log_json(event_type, message, extra=None, trace_id=None, json_dir=JSON_LOGS_DIR):
    log_entry = {
        "timestamp": datetime.utcnow().isoformat() + "Z",
        "event_type": event_type,
        "message": message,
        "trace_id": trace_id or str(uuid.uuid4()),
        "extra": extra or {}
    }
    json_log_path = os.path.join(json_dir, f"{datetime.utcnow().date()}.json")
    with open(json_log_path, "a", encoding="utf-8") as f:
        f.write(json.dumps(log_entry) + "\n")

# Dead letter logging function
def log_dead_letter(file_name, error, trace_id=None, extra=None):
    log_entry = {
        "timestamp": datetime.utcnow().isoformat() + "Z",
        "file_name": file_name,
        "error": error,
        "trace_id": trace_id or str(uuid.uuid4()),
        "extra": extra or {}
    }
    dead_letter_path = os.path.join(DEAD_LETTER_DIR, "dead_letter.log")
    with open(dead_letter_path, "a", encoding="utf-8") as f:
        f.write(json.dumps(log_entry) + "\n")

PageResult = namedtuple("PageResult", ["page_number", "markdown"])

MAX_RETRIES = int(os.getenv("MAX_RETRIES", 3))

SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_API_KEY = os.getenv("SUPABASE_API_KEY")
SUPABASE_ERROR_LOG_TABLE = "Activity_Error_Log"

# Utility function to log persistent errors to Supabase
def log_supabase_error(error_message, created_time=None):
    url = f"{SUPABASE_URL}/rest/v1/{SUPABASE_ERROR_LOG_TABLE}"
    headers = {
        "apikey": SUPABASE_API_KEY,
        "Authorization": f"Bearer {SUPABASE_API_KEY}",
        "Content-Type": "application/json",
        "Prefer": "return=minimal"
    }
    data = {
        "Activity Log/Error Message": error_message,
        "type": "error",
        "workflow name": "GCS PDF to Clean PDF",
        "Created time": created_time or datetime.utcnow().isoformat(),
        "CreatedAt": datetime.utcnow().isoformat(),
        "UpdatedAt": datetime.utcnow().isoformat(),
        "nc_order": None
    }
    try:
        response = requests.post(url, headers=headers, json=data, timeout=10)
        response.raise_for_status()
    except Exception as e:
        log_dead_letter("supabase_error", f"Failed to log to Supabase: {e}", extra={"original_error": error_message})

def split_pdf_to_pages(pdf_path: str, pdf_dir: str) -> List[str]:
    reader = PdfReader(pdf_path)
    page_files = []
    for i, page in enumerate(reader.pages):
        writer = PdfWriter()
        writer.add_page(page)
        page_path = os.path.join(pdf_dir, f"page_{i+1:04d}.pdf")
        with open(page_path, "wb") as f:
            writer.write(f)
        page_files.append(page_path)
    return page_files

def markdown_to_pdf(markdown: str, pdf_path: str, html_dir: str, page_num: int):
    html = markdown2.markdown(markdown)
    html_path = os.path.join(html_dir, f"page_{page_num:04d}.html")
    with open(html_path, "w", encoding="utf-8") as html_file:
        html_file.write(html)
    HTML(string=html).write_pdf(pdf_path)

# def markdowns_to_docx(markdown_pages: List[str], docx_path: str):
#     doc = Document()
#     for md in markdown_pages:
#         doc.add_paragraph(md)
#         doc.add_page_break()
#     doc.save(docx_path)

# def docx_to_text(docx_path):
#     doc = Document(docx_path)
#     return "\n".join([para.text for para in doc.paragraphs])

def is_valid_pdf(file_path):
    try:
        with open(file_path, 'rb') as f:
            header = f.read(5)
            if header != b'%PDF-':
                return False
            f.seek(-5, 2)
            trailer = f.read()
            if b'%%EOF' not in trailer:
                return False
        return True
    except Exception as e:
        print(f"[ERROR] Exception while validating PDF: {e}\n")
        return False

def get_pdf_page_count(pdf_path):
    try:
        reader = PdfReader(pdf_path)
        return len(reader.pages)
    except Exception as e:
        print(f"[ERROR] Could not read PDF: {e}\n")
        return 0

MAX_CONCURRENT_WORKERS = int(os.getenv("MAX_CONCURRENT_WORKERS", 8))
worker_semaphore = threading.Semaphore(MAX_CONCURRENT_WORKERS)
active_workers = 0
active_workers_lock = threading.Lock()

def get_active_worker_count():
    with active_workers_lock:
        return active_workers

# New: Global Gemini API concurrency limit
GEMINI_GLOBAL_CONCURRENCY = int(os.getenv("GEMINI_GLOBAL_CONCURRENCY", 10))
gemini_global_semaphore = threading.Semaphore(GEMINI_GLOBAL_CONCURRENCY)

MAX_CONCURRENT_FILES = int(os.getenv("MAX_CONCURRENT_FILES", 3))

def ocr_page_with_retries(pdf_path, page_number, trace_id):
    """OCR a single page with per-page retries and global Gemini API throttling."""
    for attempt in range(1, MAX_RETRIES + 1):
        try:
            with gemini_global_semaphore:
                markdown = gemini_ocr_page(pdf_path, page_number)
            return markdown
        except Exception as e:
            print(f"[ERROR][{trace_id}] OCR failed for page {page_number} (attempt {attempt}/{MAX_RETRIES}): {e}\n")
            logger.error(f"[{trace_id}] Error processing page {page_number} (attempt {attempt}/{MAX_RETRIES}): {e}")
            log_json("ocr_error", f"OCR failed for page {page_number} (attempt {attempt}/{MAX_RETRIES}): {e}", trace_id=trace_id)
            if attempt == MAX_RETRIES:
                return None
            time.sleep(2)  # brief backoff

def process_file(file_name):
    trace_id = str(uuid.uuid4())
    print(f"[START][{trace_id}] Processing file: {file_name}\n")
    logger.info(f"[{trace_id}] Processing file: {file_name}")
    log_json("start_processing", f"Processing file: {file_name}", trace_id=trace_id, extra={"trace_id": trace_id})
    retries = 0
    while retries < MAX_RETRIES:
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                pdf_dir = os.path.join(temp_dir, "pdfs")
                md_dir = os.path.join(temp_dir, "markdowns")
                html_dir = os.path.join(temp_dir, "htmls")
                os.makedirs(pdf_dir, exist_ok=True)
                os.makedirs(md_dir, exist_ok=True)
                os.makedirs(html_dir, exist_ok=True)
                print(f"[INFO][{trace_id}] Downloading from GCS to {temp_dir}\n")
                local_pdf = download_from_gcs(os.path.basename(file_name), temp_dir, trace_id=trace_id)
                print(f"[INFO][{trace_id}] Splitting PDF into pages...\n")
                page_files = split_pdf_to_pages(local_pdf, pdf_dir)
                print(f"[INFO][{trace_id}] Split into {len(page_files)} pages\n")
                results = []
                # Per-page concurrency, with per-page retries and global throttling
                with ThreadPoolExecutor(max_workers=PAGE_MAX_WORKERS) as executor:
                    futures = {executor.submit(ocr_page_with_retries, pf, i+1, trace_id): (i+1, pf) for i, pf in enumerate(page_files)}
                    for future in as_completed(futures):
                        page_number, _ = futures[future]
                        markdown = future.result()
                        if markdown is not None:
                            results.append(PageResult(page_number, markdown))
                            md_path = os.path.join(md_dir, f"page_{page_number:04d}.md")
                            with open(md_path, "w", encoding="utf-8") as md_file:
                                md_file.write(markdown)
                            print(f"[SUCCESS][{trace_id}] OCR for page {page_number} complete.\n")
                        else:
                            print(f"[ERROR][{trace_id}] OCR permanently failed for page {page_number} after {MAX_RETRIES} attempts.\n")
                            logger.error(f"[{trace_id}] OCR permanently failed for page {page_number} after {MAX_RETRIES} attempts.")
                            log_json("ocr_permanent_error", f"OCR permanently failed for page {page_number}", trace_id=trace_id)
                results.sort(key=lambda x: x.page_number)
                markdown_pages = [r.markdown for r in results]
                single_pdf_paths = []
                for i, md in enumerate(markdown_pages):
                    pdf_path = os.path.join(pdf_dir, f"ocr_page_{i+1:04d}.pdf")
                    try:
                        markdown_to_pdf(md, pdf_path, html_dir, i+1)
                        single_pdf_paths.append(pdf_path)
                        print(f"[SUCCESS][{trace_id}] Markdown to PDF for page {i+1} complete.\n")
                    except Exception as e:
                        print(f"[ERROR][{trace_id}] Markdown to PDF failed for page {i+1}: {e}\n")
                        logger.error(f"[{trace_id}] Error converting markdown to PDF for page {i+1}: {e}")
                        log_json("markdown_to_pdf_error", f"Markdown to PDF failed for page {i+1}: {e}", trace_id=trace_id)
                merged_pdf_path = os.path.join(temp_dir, "merged.pdf")
                writer = PdfWriter()
                for pdf in single_pdf_paths:
                    try:
                        reader = PdfReader(pdf)
                        writer.add_page(reader.pages[0])
                    except Exception as e:
                        print(f"[ERROR][{trace_id}] Merging page PDF failed for {pdf}: {e}\n")
                        logger.error(f"[{trace_id}] Error merging page PDF {pdf}: {e}")
                        log_json("merge_error", f"Merging page PDF failed for {pdf}: {e}", trace_id=trace_id)
                with open(merged_pdf_path, "wb") as f:
                    writer.write(f)
                print(f"[INFO][{trace_id}] Checking merged PDF at {merged_pdf_path}\n")
                merged_pdf_size = os.path.getsize(merged_pdf_path)
                print(f"[INFO][{trace_id}] Merged PDF size: {merged_pdf_size} bytes\n")
                original_page_count = get_pdf_page_count(local_pdf)
                output_page_count = get_pdf_page_count(merged_pdf_path)
                print(f"[INFO][{trace_id}] Original PDF pages: {original_page_count}, Output PDF pages: {output_page_count}\n")
                if original_page_count != output_page_count:
                    print(f"[ERROR][{trace_id}] Page count mismatch! Retrying ({retries+1}/{MAX_RETRIES})...\n")
                    logger.error(f"[{trace_id}] Page count mismatch for {file_name}. Retrying.")
                    log_json("page_count_mismatch", f"Page count mismatch for {file_name}. Retrying.", trace_id=trace_id)
                    retries += 1
                    continue
                print(f"[INFO][{trace_id}] Uploading merged PDF to GCS as {os.path.basename(file_name)}\n")
                try:
                    upload_to_gcs(merged_pdf_path, os.path.basename(file_name), trace_id=trace_id, if_generation_match=0)
                except Exception as e:
                    print(f"[ERROR][{trace_id}] Upload to GCS failed: {e}\n")
                    logger.error(f"[{trace_id}] Upload to GCS failed for {file_name}: {e}")
                    log_json("gcs_upload_error", f"Upload to GCS failed for {file_name}: {e}", trace_id=trace_id)
                    retries += 1
                    continue
                print(f"[SUCCESS][{trace_id}] Finished processing {file_name}\n")
                logger.info(f"[{trace_id}] Finished processing {file_name}")
                log_json("success", f"Finished processing {file_name}", trace_id=trace_id)
                return
        except Exception as e:
            print(f"[FATAL ERROR][{trace_id}] Processing failed for {file_name}: {e}\n")
            logger.error(f"[{trace_id}] Fatal error processing {file_name}: {e}")
            log_json("fatal_error", f"Fatal error processing {file_name}: {e}", trace_id=trace_id)
            retries += 1
    print(f"[ERROR][{trace_id}] Max retries reached for {file_name}. Skipping file.\n")
    logger.error(f"[{trace_id}] Max retries reached for {file_name}. Skipping file.")
    log_json("persistent_error", f"Max retries reached for {file_name}. Skipping file.", extra={"file_name": file_name}, trace_id=trace_id)
    log_dead_letter(file_name, "Max retries reached. Skipping file.", trace_id=trace_id)
    log_supabase_error(f"Max retries reached for {file_name}. Skipping file.", created_time=datetime.utcnow().isoformat())

# Register cleanup for graceful shutdown
_temp_dirs = []
def _register_temp_dir(path):
    _temp_dirs.append(path)
def _cleanup_temp_dirs():
    for d in _temp_dirs:
        try:
            shutil.rmtree(d, ignore_errors=True)
            logger.info(f"Cleaned up temp dir: {d}")
        except Exception as e:
            logger.error(f"Failed to clean temp dir {d}: {e}")
atexit.register(_cleanup_temp_dirs)

def _signal_handler(signum, frame):
    logger.info(f"Received signal {signum}, cleaning up and exiting...")
    _cleanup_temp_dirs()
    exit(0)
for sig in (signal.SIGINT, signal.SIGTERM):
    signal.signal(sig, _signal_handler)

# Backpressure: if too many files are queued, log a warning and sleep
MAX_QUEUE = int(os.getenv("MAX_QUEUE", 100))

def start_worker():
    in_progress = set()
    completed = set()
    pending = set()
    while True:
        try:
            with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
                future = executor.submit(list_new_files)
                try:
                    new_files = future.result(timeout=30)
                except concurrent.futures.TimeoutError:
                    print("[FATAL][WORKER] Timeout in list_new_files. Skipping this poll.")
                    time.sleep(POLL_INTERVAL)
                    continue
            # Only fetch new files if we have room to process more
            if len(in_progress) < MAX_CONCURRENT_FILES:
                for f in new_files:
                    if f not in in_progress and f not in completed and f not in pending:
                        pending.add(f)
            # Start new files if we have capacity
            while len(in_progress) < MAX_CONCURRENT_FILES and pending:
                fname = pending.pop()
                in_progress.add(fname)
                def _cb(future, fname=fname):
                    in_progress.remove(fname)
                    completed.add(fname)
                executor2 = getattr(start_worker, '_executor', None)
                if executor2 is None:
                    executor2 = concurrent.futures.ThreadPoolExecutor(max_workers=MAX_CONCURRENT_FILES)
                    setattr(start_worker, '_executor', executor2)
                future2 = executor2.submit(process_file, fname)
                future2.add_done_callback(_cb)
            time.sleep(POLL_INTERVAL)
        except Exception as e:
            print(f"[FATAL][WORKER] Exception in main worker loop: {e}\n")
            logger.error(f"[WORKER] Exception in main worker loop: {e}")
            time.sleep(POLL_INTERVAL)

# Event-driven handler for GCS notifications (to be used with Pub/Sub or HTTP trigger)
def handle_gcs_event(event_files):
    # event_files: list of file names from GCS notification
    for i in range(0, len(event_files), DOC_BATCH_SIZE):
        batch = event_files[i:i+DOC_BATCH_SIZE]
        with ThreadPoolExecutor(max_workers=PAGE_MAX_WORKERS) as executor:
            futures = [executor.submit(process_file, file) for file in batch]
            for future in as_completed(futures):
                try:
                    future.result()
                except Exception as e:
                    print(f"Error processing file: {e}")

def cleanup_old_files():
    now = datetime.utcnow()
    cutoff = now - timedelta(days=200)
    folders = [LOGS_DIR, JSON_LOGS_DIR, DEAD_LETTER_DIR, STAGING_DIR, PROCESSED_DIR]
    for folder in folders:
        if not os.path.exists(folder):
            continue
        for fname in os.listdir(folder):
            fpath = os.path.join(folder, fname)
            try:
                if os.path.isfile(fpath):
                    mtime = datetime.utcfromtimestamp(os.path.getmtime(fpath))
                    if mtime < cutoff:
                        os.remove(fpath)
                        logger.info(f"Deleted old file: {fpath}")
                elif os.path.isdir(fpath):
                    mtime = datetime.utcfromtimestamp(os.path.getmtime(fpath))
                    if mtime < cutoff:
                        shutil.rmtree(fpath, ignore_errors=True)
                        logger.info(f"Deleted old directory: {fpath}")
            except Exception as e:
                logger.error(f"Failed to delete {fpath}: {e}")

if __name__ == "__main__":
    cleanup_old_files()
    start_worker() 