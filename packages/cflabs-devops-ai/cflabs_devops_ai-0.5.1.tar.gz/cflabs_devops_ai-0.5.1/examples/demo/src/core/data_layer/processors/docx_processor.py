import hashlib
from typing import List, Dict, Any, BinaryIO
from pathlib import Path
from docx import Document
from src.base_classes.data_layer_class import DocumentProcessor
from src.core.rag_utils import chunk_document
from src.utils.logger import get_logger, time_function, log_function_call

logger = get_logger(__name__)

class DocxProcessor(DocumentProcessor):
    """Processor for DOCX documents"""
    
    def can_process(self, file_extension: str) -> bool:
        return file_extension.lower() in ['.docx', '.doc']
    
    def get_supported_extensions(self) -> List[str]:
        return ['.docx', '.doc']
    
    @time_function
    @log_function_call
    def process(self, file: BinaryIO, filename: str) -> List[Dict[str, Any]]:
        """Process DOCX file and return structured document data"""
        logger.info(f"Processing DOCX file: {filename}")
        
        try:
            # Load the document
            doc = Document(file)
            
            # Extract text from paragraphs
            content_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content_parts.append(" | ".join(row_text))
            
            content = "\n\n".join(content_parts)
            
            if not content.strip():
                logger.warning(f"Empty DOCX file: {filename}")
                return []
            
            # Generate original_uuid as hash of content
            content_hash = hashlib.sha256(content.encode()).hexdigest()
            
            # Create document
            formatted_doc = {
                "doc_id": f"{filename}_doc_1",
                "original_uuid": content_hash,
                "content": content,
                "source": filename,
                "file_type": "docx"
            }
            
            # Add chunks to the document
            chunked_doc = chunk_document(formatted_doc)
            
            logger.info(f"DOCX processing completed. Documents: 1, Total chunks: {len(chunked_doc['chunks'])}")
            return [chunked_doc]
            
        except Exception as e:
            logger.error(f"Error processing DOCX file {filename}: {str(e)}")
            raise 