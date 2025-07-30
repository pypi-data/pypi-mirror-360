"""
Tests for the Automated Labeler tool.
"""

import pytest
from pathlib import Path
from docx import Document

from scripts.tools.automated_labeler.main import AutomatedLabeler


@pytest.fixture
def tool():
    """Create an instance of the Automated Labeler tool."""
    return AutomatedLabeler()


def test_tool_initialization(tool) -> None:
    """Test tool initialization."""
    assert tool.name == "Automated Labeler"
    assert "label" in tool.description.lower()


def test_tool_run_with_valid_input(tool, sample_excel_file, sample_docx_template, temp_dir) -> None:
    """Test running the tool with valid input."""
    output_file = temp_dir / "Labels.docx"
    
    tool.run(
        input_paths=[sample_excel_file, sample_docx_template],
        output_dir=temp_dir,
        output_filename="Labels.docx"
    )
    
    # Verify output file was created
    assert output_file.exists()
    
    # Verify output file is a valid DOCX
    doc = Document(output_file)
    assert len(doc.paragraphs) > 0


def test_tool_run_without_input(tool) -> None:
    """Test running the tool without input."""
    with pytest.raises(ValueError) as exc:
        tool.run()
    assert "Both Excel file and template file paths are required" in str(exc.value)


def test_tool_run_with_missing_excel(tool, sample_docx_template, temp_dir) -> None:
    """Test running the tool with missing Excel file."""
    with pytest.raises(FileNotFoundError) as exc:
        tool.run(
            input_paths=[temp_dir / "nonexistent.xlsx", sample_docx_template],
            output_dir=temp_dir
        )
    assert "Excel file not found" in str(exc.value)


def test_tool_run_with_missing_template(tool, sample_excel_file, temp_dir) -> None:
    """Test running the tool with missing template file."""
    with pytest.raises(FileNotFoundError) as exc:
        tool.run(
            input_paths=[sample_excel_file, temp_dir / "nonexistent.docx"],
            output_dir=temp_dir
        )
    assert "Template file not found" in str(exc.value)


def test_tool_output_content(tool, sample_excel_file, sample_docx_template, temp_dir) -> None:
    """Test that the tool generates correct content."""
    output_file = temp_dir / "Labels.docx"
    
    tool.run(
        input_paths=[sample_excel_file, sample_docx_template],
        output_dir=temp_dir,
        output_filename="Labels.docx"
    )
    
    # Read the output document
    doc = Document(output_file)
    content = [p.text for p in doc.paragraphs]
    
    # Verify expected content is present
    # The exact assertions will depend on your template and data
    assert any("R001" in text for text in content)  # Research ID
    assert any("M001" in text for text in content)  # Med ID
    assert any("V1" in text for text in content)    # Visit


def test_tool_chunk_processing(tool, sample_excel_file, sample_docx_template, temp_dir) -> None:
    """Test that the tool correctly processes data in chunks."""
    output_file = temp_dir / "Labels.docx"
    
    tool.run(
        input_paths=[sample_excel_file, sample_docx_template],
        output_dir=temp_dir,
        output_filename="Labels.docx"
    )
    
    # Read the output document
    doc = Document(output_file)
    
    # Count the number of sections/pages
    # This will depend on your SETS_PER_PAGE constant and input data
    # Here we're just verifying that the document has content
    assert len(doc.paragraphs) > 0
    assert len(doc.sections) > 0


def test_tool_output_directory_creation(tool, sample_excel_file, sample_docx_template, temp_dir) -> None:
    """Test that the tool creates output directory if it doesn't exist."""
    output_dir = temp_dir / "new_output_dir"
    output_file = output_dir / "Labels.docx"
    
    tool.run(
        input_paths=[sample_excel_file, sample_docx_template],
        output_dir=output_dir,
        output_filename="Labels.docx"
    )
    
    assert output_dir.exists()
    assert output_dir.is_dir()
    assert output_file.exists()


def test_tool_empty_excel_handling(tool, temp_dir, sample_docx_template) -> None:
    """Test that the tool handles empty Excel files appropriately."""
    # Create an empty Excel file
    import pandas as pd
    empty_excel = temp_dir / "empty.xlsx"
    pd.DataFrame().to_excel(empty_excel)
    
    with pytest.raises(ValueError) as exc:
        tool.run(
            input_paths=[empty_excel, sample_docx_template],
            output_dir=temp_dir
        )
    assert "Excel file is empty" in str(exc.value) 