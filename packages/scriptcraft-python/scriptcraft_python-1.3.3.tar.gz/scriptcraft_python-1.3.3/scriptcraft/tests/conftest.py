"""
Common test configuration and fixtures.
"""

import os
import pytest
from pathlib import Path
import pandas as pd
import shutil
from typing import Generator

@pytest.fixture
def test_data_dir() -> Path:
    """Get the test data directory."""
    return Path(__file__).parent / "data"

@pytest.fixture
def temp_dir(tmp_path) -> Generator[Path, None, None]:
    """Create and clean up a temporary directory."""
    yield tmp_path
    if tmp_path.exists():
        shutil.rmtree(tmp_path)

@pytest.fixture
def temp_output_dir(tmp_path) -> Path:
    """Create a temporary output directory."""
    return tmp_path / "output"

@pytest.fixture
def sample_excel_file(test_data_dir: Path, temp_dir: Path) -> Path:
    """Create a sample Excel file for testing."""
    output_path = temp_dir / "test_input.xlsx"
    
    # Create sample data
    data = {
        'Research ID': ['R001', 'R002', 'R003'],
        'Med ID': ['M001', 'M002', 'M003'],
        'Visit': ['V1', 'V2', 'V3'],
        'AgePeriod (this is the decade of life starting at 0)': ['0-10', '11-20', '21-30'],
        'StreetNumber': ['123', '456', '789'],
        'Street Name': ['Main St', 'Oak Ave', 'Pine Rd'],
        'City/Town/Municipality': ['City1', 'City2', 'City3'],
        'State/Province': ['State1', 'State2', 'State3'],
        'Zip Code': ['12345', '67890', '11111'],
        'Country': ['USA', 'USA', 'USA'],
        'From Date': ['2020-01-01', '2020-02-01', '2020-03-01'],
        'To Date': ['2021-01-01', '2021-02-01', '2021-03-01'],
        'Comments': ['Test1', 'Test2', 'Test3']
    }
    
    df = pd.DataFrame(data)
    df.to_excel(output_path, index=True)
    return output_path

@pytest.fixture
def sample_docx_template(test_data_dir: Path, temp_dir: Path) -> Path:
    """Create a sample DOCX template for testing."""
    from docx import Document
    
    output_path = temp_dir / "template.docx"
    doc = Document()
    doc.add_paragraph("Research ID: {research_id}")
    doc.add_paragraph("Med ID: {med_id}")
    doc.add_paragraph("Visit: {visit}")
    doc.save(output_path)
    return output_path

@pytest.fixture
def sample_comparison_files(test_data_dir: Path, temp_dir: Path) -> tuple[Path, Path]:
    """Create sample files for comparison testing."""
    file1_path = temp_dir / "file1.txt"
    file2_path = temp_dir / "file2.txt"
    
    # Create two similar but slightly different files
    with open(file1_path, 'w') as f:
        f.write("Line 1\nLine 2\nLine 3\n")
    
    with open(file2_path, 'w') as f:
        f.write("Line 1\nLine 2 modified\nLine 3\n")
    
    return file1_path, file2_path 

@pytest.fixture
def sample_paths(temp_output_dir):
    """Create a sample paths dictionary."""
    return {
        "merged_data": str(temp_output_dir / "merged"),
        "qc_output": str(temp_output_dir / "qc"),
        "raw_data": str(temp_output_dir / "raw"),
        "processed_data": str(temp_output_dir / "processed")
    }