"""Document loaders for various file formats."""

import os
from pathlib import Path
from typing import Any, Dict, List, Union
from abc import ABC, abstractmethod

from pydantic import BaseModel, Field

# Remove unused magic import

from ..utils.exceptions import DocumentProcessingError
from ..utils.logging import LoggerMixin


class Document(BaseModel):
    """Represents a loaded document."""
    
    content: str = Field(description="Document content")
    metadata: Dict[str, Any] = Field(default_factory=dict, description="Document metadata")
    source: str = Field(description="Document source path or identifier")
    doc_type: str = Field(description="Document type/format")


class BaseDocumentLoader(ABC, LoggerMixin):
    """Abstract base class for document loaders."""
    
    @abstractmethod
    def load(self, source: Union[str, Path]) -> Document:
        """Load a document from the given source."""
        pass
    
    @abstractmethod
    def supports(self, file_path: Union[str, Path]) -> bool:
        """Check if this loader supports the given file."""
        pass


class TextLoader(BaseDocumentLoader):
    """Loader for plain text files."""
    
    def load(self, source: Union[str, Path]) -> Document:
        """Load a text file."""
        try:
            with open(source, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return Document(
                content=content,
                metadata={
                    "file_size": os.path.getsize(source),
                    "encoding": "utf-8"
                },
                source=str(source),
                doc_type="text"
            )
        except Exception as e:
            raise DocumentProcessingError(f"Failed to load text file: {e}", document_path=str(source))
    
    def supports(self, file_path: Union[str, Path]) -> bool:
        """Check if file is a text file."""
        return str(file_path).lower().endswith(('.txt', '.md', '.markdown'))


class PDFLoader(BaseDocumentLoader):
    """Loader for PDF files."""
    
    def __init__(self, extract_images: bool = False):
        self.extract_images = extract_images
    
    def load(self, source: Union[str, Path]) -> Document:
        """Load a PDF file."""
        try:
            import pypdf
            
            with open(source, 'rb') as f:
                reader = pypdf.PdfReader(f)
                content = ""
                
                for page in reader.pages:
                    content += page.extract_text() + "\n"
            
            metadata: Dict[str, Any] = {
                "num_pages": len(reader.pages),
                "file_size": os.path.getsize(source)
            }

            # Extract metadata from PDF
            if reader.metadata:
                pdf_metadata: Dict[str, Any] = {
                    "title": str(reader.metadata.get("/Title", "")),
                    "author": str(reader.metadata.get("/Author", "")),
                    "subject": str(reader.metadata.get("/Subject", "")),
                    "creator": str(reader.metadata.get("/Creator", ""))
                }
                metadata.update(pdf_metadata)
            
            return Document(
                content=content.strip(),
                metadata=metadata,
                source=str(source),
                doc_type="pdf"
            )
        except Exception as e:
            raise DocumentProcessingError(f"Failed to load PDF file: {e}", document_path=str(source))
    
    def supports(self, file_path: Union[str, Path]) -> bool:
        """Check if file is a PDF."""
        return str(file_path).lower().endswith('.pdf')


class DOCXLoader(BaseDocumentLoader):
    """Loader for DOCX files."""
    
    def load(self, source: Union[str, Path]) -> Document:
        """Load a DOCX file."""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(str(source))
            content = ""
            
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            metadata: Dict[str, Any] = {
                "file_size": os.path.getsize(source),
                "num_paragraphs": len(doc.paragraphs)
            }

            # Extract core properties
            if doc.core_properties:
                docx_metadata: Dict[str, Any] = {
                    "title": str(doc.core_properties.title or ""),
                    "author": str(doc.core_properties.author or ""),
                    "subject": str(doc.core_properties.subject or ""),
                    "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                    "modified": str(doc.core_properties.modified) if doc.core_properties.modified else ""
                }
                metadata.update(docx_metadata)
            
            return Document(
                content=content.strip(),
                metadata=metadata,
                source=str(source),
                doc_type="docx"
            )
        except Exception as e:
            raise DocumentProcessingError(f"Failed to load DOCX file: {e}", document_path=str(source))
    
    def supports(self, file_path: Union[str, Path]) -> bool:
        """Check if file is a DOCX."""
        return str(file_path).lower().endswith('.docx')


class HTMLLoader(BaseDocumentLoader):
    """Loader for HTML files."""
    
    def __init__(self, parser: str = "html.parser"):
        self.parser = parser
    
    def load(self, source: Union[str, Path]) -> Document:
        """Load an HTML file."""
        try:
            from bs4 import BeautifulSoup
            
            with open(source, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, self.parser)
            
            # Extract text content
            content = soup.get_text(separator='\n', strip=True)
            
            # Extract metadata
            metadata = {
                "file_size": os.path.getsize(source),
                "title": soup.title.string if soup.title else "",
            }
            
            # Extract meta tags
            meta_tags = soup.find_all('meta')
            for tag in meta_tags:
                if tag.get('name'):
                    metadata[f"meta_{tag.get('name')}"] = tag.get('content', '')
            
            return Document(
                content=content,
                metadata=metadata,
                source=str(source),
                doc_type="html"
            )
        except Exception as e:
            raise DocumentProcessingError(f"Failed to load HTML file: {e}", document_path=str(source))
    
    def supports(self, file_path: Union[str, Path]) -> bool:
        """Check if file is HTML."""
        return str(file_path).lower().endswith(('.html', '.htm'))


class DocumentLoader(LoggerMixin):
    """Main document loader that delegates to specific loaders."""
    
    def __init__(self):
        self.loaders = [
            TextLoader(),
            PDFLoader(),
            DOCXLoader(),
            HTMLLoader(),
        ]
    
    def load_file(self, file_path: Union[str, Path]) -> Document:
        """Load a single file."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise DocumentProcessingError(f"File not found: {file_path}")
        
        # Find appropriate loader
        for loader in self.loaders:
            if loader.supports(file_path):
                self.logger.debug(f"Loading {file_path} with {loader.__class__.__name__}")
                return loader.load(file_path)
        
        # Fallback to text loader for unknown types
        self.logger.warning(f"No specific loader found for {file_path}, using text loader")
        return TextLoader().load(file_path)
    
    def load_directory(self, directory_path: Union[str, Path], recursive: bool = True) -> List[Document]:
        """Load all supported files from a directory."""
        directory_path = Path(directory_path)
        
        if not directory_path.exists():
            raise DocumentProcessingError(f"Directory not found: {directory_path}")
        
        documents = []
        pattern = "**/*" if recursive else "*"
        
        for file_path in directory_path.glob(pattern):
            if file_path.is_file():
                try:
                    document = self.load_file(file_path)
                    documents.append(document)
                except DocumentProcessingError as e:
                    self.logger.warning(f"Failed to load {file_path}: {e}")
                    continue
        
        self.logger.info(f"Loaded {len(documents)} documents from {directory_path}")
        return documents
    
    def load_multiple(self, file_paths: List[Union[str, Path]]) -> List[Document]:
        """Load multiple files."""
        documents = []
        
        for file_path in file_paths:
            try:
                document = self.load_file(file_path)
                documents.append(document)
            except DocumentProcessingError as e:
                self.logger.warning(f"Failed to load {file_path}: {e}")
                continue
        
        return documents
