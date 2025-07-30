import io
import os
import re
import shutil
import sys
from configparser import ConfigParser
from datetime import datetime

from PySide6 import QtCore
from PySide6.QtCore import QEventLoop, QTimer
from PySide6.QtWidgets import QApplication, QMainWindow, QFileDialog, QInputDialog
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.opc.exceptions import PackageNotFoundError
from openpyxl.reader.excel import load_workbook
from win32com.client import DispatchEx

from form import Ui_MainWindow


class Project(object):
    p_comname = ''  # 公司名称
    p_order = ''  # 序号
    p_name = ''  # 项目名
    p_start = ''
    p_end = ''
    p_cost = ''
    p_people = ''  # 人数
    p_owner = ''  # 项目负责人
    p_rnd = ''  # 研发人员
    p_money = ''  # 总预算
    pat_list = ''  # 知识产权名称
    ip_list = ''  # IP
    p_product = ''  # PS


class Patent(object):
    p_patname = ''  # 知识产权名称
    p_order = ''  # 序号
    p_class = ''  # 类别
    p_no = ''  # 专利/登记号
    p_got = ''  # 获得方式
    p_date = ''  # 授权日期

class CheckR(object):
    """文档检查结果"""

    def __init__(self, match=0, unmatch=0):
        self.match = match
        self.unmatch = unmatch

    def __add__(self, other):
        match = self.match + other.match
        unmatch = self.unmatch + other.unmatch
        return CheckR(match, unmatch)


class EmittingStr(QtCore.QObject):
    textWritten = QtCore.Signal(str)

    def write(self, text):
        self.textWritten.emit(str(text))
        loop = QEventLoop()
        QTimer.singleShot(100, loop.quit)
        loop.exec()
        QApplication.processEvents()


def format_date(date_str, index):
    if isinstance(date_str, datetime):
        return date_str.strftime('%Y-%m-%d')
    else:
        try:
            tmpdate = datetime.strptime(date_str, '%Y-%m-%d')
            return tmpdate.strftime('%Y-%m-%d')
        except (ValueError, TypeError) as error:
            print(f"序号{index} 日期格式错误:{date_str}")
            return date_str


class MainWindow(Ui_MainWindow, QMainWindow):
    workdir = ''
    file_prj = ''  # 立项报告汇总表
    file_pat = ''  # 知识产权汇总表
    file_pat_upload = ''  # 知识产权表_上传
    pat_list = []  # 专利对象列表
    pat_dict = {}  # 专利名字->序号字典   TODO 合并到dict2里面去
    pat_dict2 = {}  # 专利名字->专利对象字典
    arr_prj = []  # 项目数组
    com_name = ''

    def __init__(self):
        super(MainWindow, self).__init__()
        self.setupUi(self)
        sys.stdout = EmittingStr()
        sys.stdout.textWritten.connect(self.outputwritten)
        sys.stderr = EmittingStr()
        sys.stderr.textWritten.connect(self.outputwritten)

        self.actionSelect_Dir.triggered.connect(self.set_docurl)
        self.actioncheck.triggered.connect(lambda: self.checkpatent(True))
        self.actionreplace.triggered.connect(lambda: self.replaceprj(True))
        self.actioncheckall.triggered.connect(self.checkall)
        self.actionsearchall.triggered.connect(self.searchall)
        self.actionpat_table.triggered.connect(self.prepare_pat_table)
        self.actionremove_yellow.triggered.connect(self.remove_highlight)
        self.actionrd_summary.triggered.connect(self.prepare_rdsummary_table)

        self.config = ConfigParser()
        try:
            self.config.read('config.ini', encoding='UTF-8')
            self.workdir = self.config['config']['lasting']
        except KeyError:
            self.config.add_section('config')
        try:
            if self.workdir != '':
                self.onchangeworkdir()
        except FileNotFoundError as e:
            self.textEdit.append('路径错误：' + e.filename)
            self.workdir = ''
        self.lineEdit.setText(self.workdir)

    def outputwritten(self, text):
        self.textEdit.append(text.strip())

    def set_docurl(self):
        # 重新选择输入和输出目录时，进度条设置为0，文本框的内容置空
        tempdir = QFileDialog.getExistingDirectory(self, "选中项目所在目录", r"")
        if tempdir != '':
            self.workdir = tempdir
            self.onchangeworkdir()
            with open('config.ini', 'w', encoding='utf-8') as file:
                self.config['config']['lasting'] = self.workdir
                self.config.write(file)  # 数据写入配置文件

    def onchangeworkdir(self):
        self.lineEdit.setText(self.workdir)
        self.textEdit.clear()
        for file_sum in os.listdir(self.workdir):
            if file_sum.endswith('立项报告汇总表.xlsx') and not file_sum.startswith('~$'):
                self.file_prj = self.workdir + '/' + file_sum
            if file_sum.endswith('知识产权汇总表.xlsx') and not file_sum.startswith('~$'):
                self.file_pat = self.workdir + '/' + file_sum
            if file_sum.endswith('1.知识产权表（参与本次创新能力知识产权评价，汇总信息只统计此列表中的知识产权.xlsx') and not file_sum.startswith('~$'):
                self.file_pat_upload = self.workdir + '/' + file_sum
        if self.file_prj == '':
            self.textEdit.append('没找到：' + '立项报告汇总表.xlsx')
        if self.file_pat == '':
            self.textEdit.append('没找到：' + '知识产权汇总表.xlsx')
        if self.file_pat_upload == '':
            self.textEdit.append('没找到：' + '1.知识产权表（参与本次创新能力知识产权评价，汇总信息只统计此列表中的知识产权.xlsx')
        # self.update_data()

    def replaceprj(self, modify=False):
        if modify:
            self.textEdit.setText('')
            self.update_data()

        for project in self.arr_prj:
            # self.textEdit.append(project.p_order + ' 项目开始处理...')
            checkr = CheckR()  # 已匹配条目数
            doc_name = ''
            try:
                doc_name = self.workdir + '/RD' + project.p_order + '_' + project.p_name + '.docx'
                old_doc_name = self.workdir + '/RD' + project.p_order + project.p_name + '.docx'
                if os.path.isfile(old_doc_name) and not os.path.isfile(doc_name):
                    # 如果存在，重命名文件
                    new_file_path = 'example_renamed.txt'
                    shutil.move(old_doc_name, doc_name)
                    print(f'File renamed to {new_file_path}')

                document = Document(doc_name)
                # debug_doc(document)
                checkr = checkr + self.replace_comname(document, project)
                checkr = checkr + self.first_table(document, project)
                checkr = checkr + self.start_time(document, project)
                checkr = checkr + self.second_table(document, project)
                checkr = checkr + self.third_table(document, project)

                checkr = checkr + self.checkpat2(document, project)

                if modify:
                    document.save(doc_name)
            except PackageNotFoundError:
                self.textEdit.append('Error打开文件错误：' + doc_name)
            except PermissionError:
                self.textEdit.append('Error 保存文件错误，可能是文件已被打开：' + doc_name)

            prompt = project.p_order + ' 项目检查完成。 <font color="green"><b>' + str(
                checkr.match) + ' </b></font> 项匹配。'
            if checkr.unmatch > 0:
                prompt = prompt + ' <font color="red"><b>' + str(checkr.unmatch) + ' </b></font> 项不匹配。'
            self.textEdit.append(prompt)

    def checkpatent(self, modify=False):
        if modify:
            self.textEdit.setText('')
            self.update_data()

        match = 0  # 已匹配条目数
        unmatch = 0  # 已匹配条目数
        changed = False
        wb = load_workbook(self.file_prj)
        ws = wb.active
        max_row_num = ws.max_row
        range_cell = ws[f'A3:P{max_row_num}']
        i: int = 0
        for r in range_cell:
            if r[11].value is None:
                break

            pat_name = str(r[11].value).strip()
            # 检查专利IP号
            if pat_name == '无':
                if r[14].value != '无':
                    self.textEdit.append(self.arr_prj[i].p_order + ' 项目: ' + str(r[14].value) + ' ===> ' + '无')
                    r[14].value = pat_name
                    changed |= True
                    unmatch = unmatch + 1
                else:
                    match = match + 1
            else:
                lst = pat_name.splitlines()
                rep = [self.pat_dict[x] if x in self.pat_dict else x for x in lst]  # 神啊，不懂了
                for element in rep:
                    if re.search(r'^\d\d$', element) is None:
                        self.textEdit.append('专利IP错误：' + element)
                rep = map(lambda e: 'IP' + e, rep)
                new_ip = ','.join(rep)
                if r[14].value != new_ip:
                    self.textEdit.append(self.arr_prj[i].p_order + ' 项目: ' + str(r[14].value) + ' ===> ' + new_ip)
                    r[14].value = new_ip
                    changed |= True
                    unmatch = unmatch + 1
                else:
                    match = match + 1

            # 检查和替换项目人数
            rnd_name = str(r[8].value).strip()
            if rnd_name != '' and not rnd_name.endswith('等'):
                real_num = len(rnd_name.split('、'))
                real_num = real_num + 1
                if real_num != int(r[6].value):
                    self.textEdit.append(self.arr_prj[i].p_order + ' 项目: ' + '研发人员名字数量不匹配:' + rnd_name)
                    self.textEdit.append('研发人数: ' + str(r[6].value) + ' ===> ' + str(real_num))
                    r[6].value = real_num
                    changed |= True
                    unmatch = unmatch + 1
                else:
                    match = match + 1
            i = i + 1
        try:
            if changed and modify:
                wb.save(self.file_prj)
                xl_app = DispatchEx("Excel.Application")
                xl_app.Visible = False
                xl_app.DisplayAlerts = False
                xl_book = xl_app.Workbooks.Open(self.file_prj)
                xl_book.Saved = False
                xl_book.Close(True)
                xl_app.Quit()
        except PermissionError:
            self.textEdit.append('写文件失败，关闭其他占用该文件的程序.' + self.file_prj)
        wb.close()

        prompt = '立项表IP更新完成。 <font color="green"><b>' + str(match) + ' </b></font> 项条目匹配。'
        if unmatch != 0:
            prompt = prompt + ' <font color="red"><b>' + str(unmatch) + ' </b></font> 项条目不匹配。'
        self.textEdit.append(prompt)

    def checkpat2(self, doc, prj):
        match = unmatch = 0
        lst = prj.pat_list.splitlines()
        for pat_name in lst:
            if pat_name in self.pat_dict2:
                pat = self.pat_dict2[pat_name]
                check_str = '知识产权类型错误时默认字符串'
                match pat.p_class:
                    case '软件著作权' | '集成电路布图设计专有权':
                        check_str = pat.p_name + '，登记号：' + pat.p_no
                    case '实用新型专利' | '外观设计专利':
                        check_str = pat.p_name + '，授权号：' + pat.p_no
                    case '发明专利（非国防专利）' | '发明专利（国防专利）':
                        check_str = pat.p_name + '，授权号：' + pat.p_no
                    case _:
                        self.textEdit.append('知识产权类型错误：' + pat.p_name + '：' + pat.p_class)
                found = False
                for i, para in enumerate(doc.tables[2].rows[4].cells[0].paragraphs):
                    if check_str in para.text:
                        found |= True
                        match = match + 1
                if not found:
                    unmatch = unmatch + 1
                    self.textEdit.append('全文找不到：' + check_str)
            elif pat_name != "无":
                unmatch = unmatch + 1
                self.textEdit.append('没有找到知识产权：' + pat_name)

        #  TODO  统计软件、发明件数， 检查总数
        # if pat1 != 0:
        #    check_replace(self, doc.tables[2].rows[4].cells[0].paragraphs, regex, dst)

        return CheckR(match, unmatch)

    def update_data(self):
        self.com_name = ''
        self.pat_list.clear()
        self.pat_dict.clear()
        self.pat_dict2.clear()

        with open(self.file_pat, "rb") as f:
            in_mem_file = io.BytesIO(f.read())
        wb = load_workbook(in_mem_file, read_only=True, data_only=True)
        ws = wb.active
        if ws is not None:
            max_row_num = ws.max_row
            range_cell = ws[f'A3:F{max_row_num}']
            for r in range_cell:
                if r[0].value is None:
                    break
                pat = Patent()
                pat.p_order = str(r[0].value).strip().zfill(2)
                pat.p_name = str(r[1].value).strip()
                pat.p_class = str(r[2].value).strip()
                pat.p_no = str(r[3].value).strip()
                pat.p_date = format_date(r[4].value, pat.p_order)
                pat.p_got = str(r[5].value).strip()

                self.pat_list.append(pat)
                self.pat_dict[pat.p_name] = pat.p_order
                self.pat_dict2[pat.p_name] = pat
        else:
            self.textEdit.append('Error ' + self.file_pat + ' 文件格式错误。')
        wb.close()

        self.arr_prj.clear()
        with open(self.file_prj, "rb") as f:
            in_mem_file = io.BytesIO(f.read())
        wb = load_workbook(in_mem_file, read_only=True, data_only=True)
        ws = wb.active
        if ws is not None:
            if str(ws['A1'].value).find(u'公司') != -1:
                self.com_name = str(ws['A1'].value).split("公司")[0] + '公司'
            else:
                self.textEdit.append("Error: 找不到 公司名")
            max_row_num = ws.max_row
            range_cell = ws[f'A3:P{max_row_num}']
            for r in range_cell:
                if r[0].value is None:
                    break
                project = Project()
                project.p_comname = self.com_name
                project.p_order = str(r[0].value).strip().zfill(2)
                project.p_name = str(r[1].value).strip()  # 项目名称
                project.p_start = format_date(r[2].value, project.p_order)
                project.p_end = format_date(r[3].value, project.p_order)
                project.p_cost = str(r[5].value).strip()
                project.p_people = str(r[6].value).strip()  # 人数
                project.p_owner = str(r[7].value).strip()  # 项目负责人
                project.p_rnd = str(r[8].value).strip()  # 研发人员
                project.p_money = str(r[9].value).strip()  # 总预算
                project.pat_list = str(r[11].value).strip()  # 知识产权名称
                project.ip_list = str(r[14].value).strip()  # IP
                rst = re.findall(r"(PS[0-9][0-9])：.*", str(r[15].value).strip())
                project.p_product = ','.join(rst)
                self.arr_prj.append(project)
        else:
            self.textEdit.append('Error ' + self.file_prj + ' 文件格式错误。')
        wb.close()
        if len(self.arr_prj) == 0:
            self.textEdit.append("Error: 立项汇总表错误，使用Excel重新保存.")

    def checkall(self):
        self.textEdit.setText('')
        self.update_data()

        self.checkpatent(False)
        self.replaceprj(False)

    def searchall(self):
        self.textEdit.setText('')
        self.update_data()
        text, ok = QInputDialog.getText(self, "全项目查找",
                                        "查找内容:")
        if ok and text:
            for project in self.arr_prj:
                # self.textEdit.append(project.p_order + ' 项目开始处理...')
                checkr = CheckR()  # 已匹配条目数
                doc_name = ''
                try:
                    doc_name = self.workdir + '/RD' + project.p_order + '_' + project.p_name + '.docx'
                    document = Document(doc_name)
                    checkr = checkr + self.findindoc(document, text)
                except PackageNotFoundError:
                    self.textEdit.append('Error打开文件错误：' + doc_name)
                except PermissionError:
                    self.textEdit.append('Error保存文件错误，可能是文件已被打开：' + doc_name)

                prompt = project.p_order + ' 项目检查完成。 <font color="green"><b>' + str(
                    checkr.match) + ' </b></font> 项匹配。'
                if checkr.unmatch > 0:
                    prompt = prompt + ' <font color="red"><b>' + str(checkr.unmatch) + ' </b></font> 项不匹配。'
                self.textEdit.append(prompt)

    def remove_highlight(self):
        self.textEdit.setText('')
        self.update_data()
        for project in self.arr_prj:
            # self.textEdit.append(project.p_order + ' 项目开始处理...')
            checkr = CheckR()  # 已匹配条目数
            doc_name = ''
            try:
                doc_name = self.workdir + '/RD' + project.p_order + '_' + project.p_name + '.docx'
                document = Document(doc_name)
                checkr = checkr + self.remove_hl_indoc(document)
                if checkr.match > 0:
                    prompt = project.p_order + ' 项目检查完成。 <font color="green"><b>' + str(
                        checkr.match) + ' </b></font> 项黄底字体已替换。'
                    document.save(doc_name)
                else:
                    prompt = project.p_order + ' 项目检查完成。 没有黄底字体。'
                self.textEdit.append(prompt)
            except PackageNotFoundError:
                self.textEdit.append('Error打开文件错误：' + doc_name)
            except PermissionError:
                self.textEdit.append('Error保存文件错误，可能是文件已被打开：' + doc_name)

    @staticmethod
    def clear_runs(runs):
        for i, run in enumerate(runs):
            if i > 0:
                run.clear()
        return runs

    @staticmethod
    def clear_paras(paras):
        para = paras[0]
        paras.clear()
        paras.append(para)
        return paras

    def debug_doc(self, doc):
        for i, sect in enumerate(doc.sections):
            for j, para in enumerate(sect.header.paragraphs):
                self.textEdit.append(f'Sec.{i} Para.{j} : ', para.text, sep='')
                for k, run in enumerate(para.runs):
                    self.textEdit.append(f'Sec.{i} Para.{j} Run{k}: ', run.text, sep='')
        for i, para in enumerate(doc.paragraphs):
            self.textEdit.append(f'Para.{i} : ', para.text, sep='')
            for j, run in enumerate(para.runs):
                self.textEdit.append(f'Para.{i} Run{j}: ', run.text, sep='')
        for i, table in enumerate(doc.tables):
            for j, row in enumerate(table.rows):
                for k, cell in enumerate(row.cells):
                    for l, para in enumerate(cell.paragraphs):
                        self.textEdit.append(f'Table.{i} Row.{j} Cell{k} Para.{l} : ', para.text, sep='')
                        # for j, run in enumerate(para.runs):
                        #     self.textEdit.append(f'Para.{i} Run{j}: ', run.text, sep='')

    def findindoc(self, doc, keyword):
        match = 0
        for i, sect in enumerate(doc.sections):
            for j, para in enumerate(sect.header.paragraphs):
                if keyword in para.text:
                    match = match + 1
                    newstr = para.text.replace(keyword, f'<font color="red"><b>{keyword}</b></font>')
                    self.textEdit.append(f'节.{i} 段.{j} : {newstr}')
                # for k, run in enumerate(para.runs):
                #     self.textEdit.append(f'Sec.{i} Para.{j} Run{k}: ', run.text, sep='')
        for i, para in enumerate(doc.paragraphs):
            if keyword in para.text:
                match = match + 1
                newstr = para.text.replace(keyword, f'<font color="red"><b>{keyword}</b></font>')
                self.textEdit.append(f'段.{i} : {newstr}')
                # for j, run in enumerate(para.runs):
                #     self.textEdit.append(f'段.{i} Run{j}: ', run.text, sep='')
        for i, table in enumerate(doc.tables):
            for j, row in enumerate(table.rows):
                for k, cell in enumerate(row.cells):
                    for l, para in enumerate(cell.paragraphs):
                        if keyword in para.text:
                            match = match + 1
                            newstr = para.text.replace(keyword, f'<font color="red"><b>{keyword}</b></font>')
                            self.textEdit.append(f'表.{i} 行.{j} 列{k} 段.{l} : {newstr}')
        return CheckR(match, 0)

    def remove_hl_indoc(self, doc):
        match = 0
        for i, sect in enumerate(doc.sections):
            for j, para in enumerate(sect.header.paragraphs):
                for k, run in enumerate(para.runs):
                    if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                        self.textEdit.append(f'Highlight已去除:  {run.text}')
                        run.font.highlight_color = None
                        match = match + 1
                    elif run.font.highlight_color is not None:
                        self.textEdit.append(f'其它Highlight颜色，未处理:  {run.text}')
        for i, para in enumerate(doc.paragraphs):
            for k, run in enumerate(para.runs):
                if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                    self.textEdit.append(f'Highlight已去除:  {run.text}')
                    run.font.highlight_color = None
                    match = match + 1
                elif run.font.highlight_color is not None:
                    self.textEdit.append(f'其它Highlight颜色，未处理:  {run.text}')
        for i, table in enumerate(doc.tables):
            for j, row in enumerate(table.rows):
                for k, cell in enumerate(row.cells):
                    for l, para in enumerate(cell.paragraphs):
                        for m, run in enumerate(para.runs):
                            if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                                self.textEdit.append(f'Highlight已去除:  {run.text}')
                                run.font.highlight_color = None
                                match = match + 1
                            elif run.font.highlight_color is not None:
                                self.textEdit.append(f'其它Highlight颜色，未处理:  {run.text}')
        return CheckR(match, 0)

    def replace_comname(self, doc, prj):
        match = unmatch = 0
        para = doc.sections[0].header.paragraphs[0]
        oldname = para.text.strip()
        if oldname == '':
            unmatch = unmatch + 1
            self.textEdit.append('没有在页眉检测到公司名称')
            return CheckR(match, unmatch)

        if prj.p_comname in para.text:
            match = match + 1
        else:
            unmatch = unmatch + 1
            self.textEdit.append(oldname + ' ===> ' + prj.p_comname)
            # TODO 这个地方还有个问题， 先用临时凑合一下
            # found = False
            # for i, run in enumerate(para.runs):
            #     if found:
            #         run.clear()
            #     if run.text != '':
            #         run.text = '\t\t' + prj.p_comname
            #         found = True
            for i, run in enumerate(para.runs):
                if run.text == oldname:
                    run.text = prj.p_comname
                    self.textEdit.append(oldname + ' ===> ' + prj.p_comname)
                    match = match + 1
        checkr = CheckR(match, unmatch) + self.check_replace(doc.paragraphs, oldname, prj.p_comname)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    checkr = checkr + self.check_replace(cell.paragraphs, oldname, prj.p_comname)

        # 检查其他可能的未替换的内容，例如去掉深圳，有限公司等
        # TODO

        return checkr

    def first_table(self, doc, prj):
        checkr = self.check_replace_para(doc.tables[0].rows[0].cells[1].paragraphs[0], prj.p_name)
        checkr = checkr + self.check_replace_para(doc.tables[0].rows[1].cells[1].paragraphs[0],
                                                 prj.p_start[0:4] + 'RD' + prj.p_order)
        checkr = checkr + self.check_replace_para(doc.tables[0].rows[2].cells[1].paragraphs[0], prj.p_owner)
        checkr = checkr + self.check_replace_para(doc.tables[0].rows[3].cells[1].paragraphs[0],
                                                  prj.p_start + '至' + prj.p_end)
        return checkr

    def start_time(self, doc, prj):
        return self.check_replace(doc.paragraphs, r'申请立项时间：\d{2,4}[-/]\d{1,2}[-/]\d{1,2}',
                                  '申请立项时间：' + prj.p_start)

    def second_table(self, doc, prj):
        checkr = self.check_replace_para(doc.tables[1].rows[0].cells[1].paragraphs[0], prj.p_name)

        checkr = checkr + self.check_replace(doc.tables[1].rows[1].cells[1].paragraphs,
                                             '项目团队由(.*)人组成，项目实施周期为(.*)个月。',
                                             '项目团队由' + prj.p_people + '人组成，项目实施周期为' + prj.p_cost + '个月。')
        if len(doc.tables[1].rows[6].cells[1].paragraphs) > 1:
            self.textEdit.append('<font color="red"><b>' + '项目进度安排->工作进度 内容为多段，需要删除多余的段落. ' +
                                 ' </b></font>' )
        checkr = checkr + self.check_replace_paras(doc.tables[1].rows[6].cells[1].paragraphs,
                                             prj.p_start + '\n至\n' + prj.p_end)
        checkr = checkr + self.check_replace(doc.tables[1].rows[7].cells[1].paragraphs,
                                             '项目总资金预算.*万元', '项目总资金预算' + prj.p_money + '万元')

        checkr = checkr + self.check_replace(doc.tables[1].rows[8].cells[1].paragraphs,
                                             '项目总人数：.*人', '项目总人数：' + prj.p_people + '人')
        if prj.p_owner != 'None':
            checkr = checkr + self.check_replace(doc.tables[1].rows[8].cells[1].paragraphs, '项目负责人：.*',
                                                 '项目负责人：' + prj.p_owner)
        if prj.p_rnd != 'None':
            checkr = checkr + self.check_replace(doc.tables[1].rows[8].cells[1].paragraphs, '研发成员：.*',
                                                 '研发成员：' + prj.p_rnd)
        checkr = checkr + self.check_replace(doc.tables[1].rows[9].cells[1].paragraphs,
                                             r'\d{2,4}[-/]\d{1,2}[-/]\d{1,2}',
                                             prj.p_start)
        return checkr

    def check_replace(self, paras, regex, dst):
        """
        在多个段落中查找和匹配，只进行一次匹配
        """
        match = unmatch = 0
        for i, para in enumerate(paras):
            result = re.search(regex, para.text)
            if result is not None:
                if result.group() != dst and len(para.runs) > 0:
                    unmatch = unmatch + 1
                    self.textEdit.append(str(result.group()) + ' ===> ' + dst)
                    para.runs[0].text = re.sub(regex, dst,
                                               para.text)
                    self.clear_runs(para.runs)
                else:
                    match = match + 1
                break  # 只替换一次就够用
        return CheckR(match, unmatch)

    def check_replace_para(self, para, dst):
        """
        检查和替换掉所传入的 para
        """
        match = unmatch = 0
        if para.text != dst:
            unmatch = unmatch + 1
            self.textEdit.append(para.text + ' ===> ' + dst)
            if len(para.runs) == 0:
                para.add_run('')
            para.runs[0].text = dst
            self.clear_runs(para.runs)
        else:
            match = match + 1
        return CheckR(match, unmatch)

    def check_replace_paras(self, paras, dst):
        """
        检查和替换掉所传入的 paras 对象，确保整体内容匹配

        TOTO:表格中传入的paras对象实际为复制的对象，无法修改原有对象，后续看有没有改进
        https://github.com/python-openxml/python-docx/issues/33
        """
        match = unmatch = 0
        org_text = ""
        for i, para in enumerate(paras):
            org_text +=  para.text
        if org_text.replace("\n","") != dst.replace("\n",""):
            unmatch = unmatch + 1
            self.textEdit.append(org_text + ' ===> ' + dst)

            # 临时方案, 修改 paras对象无法传出，后续看是否有其它方案
            for i, para in enumerate(paras):
                for j, run in enumerate(para.runs):
                    run.clear()

            if len(paras[0].runs) == 0:
                paras[0].add_run('dst')
            else:
                paras[0].runs[0].text = dst
            self.clear_paras(paras)
            self.clear_runs(paras[0].runs)
        else:
            match = match + 1
        return CheckR(match, unmatch)

    def third_table(self, doc, prj):
        """
        项目验收报告
        """
        checkr = self.check_replace_para(doc.tables[2].rows[0].cells[1].paragraphs[0], prj.p_name)
        checkr = checkr + self.check_replace_para(doc.tables[2].rows[1].cells[1].paragraphs[0], prj.p_end)
        checkr = checkr + self.check_replace_para(doc.tables[2].rows[2].cells[1].paragraphs[0],
                                             prj.p_start + '至' + prj.p_end)
        checkr = checkr + self.check_replace_para(doc.tables[2].rows[3].cells[1].paragraphs[0], prj.p_owner)
        return checkr

    def get_prj_init(self, project):
        """
        获取项目 立项目的、意义及组织实施方式//核心技术及创新点//取得的阶段性成果
        """
        rst = False
        tmp1 = tmp2 = doc_name = ''
        try:
            doc_name = self.workdir + '/RD' + project.p_order + '_' + project.p_name + '.docx'
            doc = Document(doc_name)
            paras = doc.tables[1].rows[1].cells[1].paragraphs
            for i, para in enumerate(paras):
                tmp1 = tmp1 + para.text
            tmp1 = tmp1.replace("1、立项目的及意义", "")
            tmp1 = tmp1.replace("2、组织实施方式", "")
            tmp1 = tmp1.split("主要任务分配：", 1)[0]

            paras = doc.tables[2].rows[4].cells[0].paragraphs
            for i, para in enumerate(paras):
                tmp2 = tmp2 + para.text
            tmp2 = tmp2.split("1、取得的阶段性成果", 1)[1]
            tmpcg_tmp= tmp2.split("。", 1)[0] + "。"
            'tmpcg = tmp2.split("2、成效", 1)[0]'
            tmp2 = tmp2.split("2、成效", 1)[1]
            tmpcx = tmp2.split("3、项目验收情况", 1)[0]
            tmpcg = tmpcg_tmp + tmpcx
            rst = True
        except PackageNotFoundError:
            self.textEdit.append('Error打开文件错误：' + doc_name)
        except PermissionError:
            self.textEdit.append('Error 保存文件错误，可能是文件已被打开：' + doc_name)
        return rst, tmp1, tmpcg, tmpcx

    def get_prj_core(self, project):
        """
        获取项目核心技术和创新点
        """
        tmp = doc_name = ''
        try:
            doc_name = self.workdir + '/RD' + project.p_order + '_' + project.p_name + '.docx'
            doc = Document(doc_name)
            paras = doc.tables[1].rows[2].cells[1].paragraphs
            for i, para in enumerate(paras):
                tmp = tmp + para.text + '\n'
            rst = re.search(r"1、核心技术\n(.*)2、创新点\n(.*)", tmp, re.S | re.M)
            if rst is not None:
                return True, rst.group(1), rst.group(2), tmp
            else:
                self.textEdit.append(doc_name + ': 没有找到 核心技术 or 创新点')
                return False, '', '',''
        except PackageNotFoundError:
            self.textEdit.append('Error打开文件错误：' + doc_name)
        except PermissionError:
            self.textEdit.append('Error 保存文件错误，可能是文件已被打开：' + doc_name)
        return False, '', '', ''

    def prepare_pat_table(self):
        """
        根据专利汇总表准备上传的表格
        """
        self.textEdit.setText('')
        self.update_data()
        wb = load_workbook(self.file_pat_upload)
        ws = wb.active

        row = 2
        for pat in self.pat_list:
            ws.cell(row, column=1, value="IP" + pat.p_order)
            ws.cell(row, column=2, value=pat.p_name)
            ws.cell(row, column=3, value=pat.p_class)
            ws.cell(row, column=4, value=pat.p_got)
            ws.cell(row, column=5, value=pat.p_no)
            ws.cell(row, column=6, value=pat.p_date)
            ws.cell(row, column=7, value=self.com_name)
            tmp1 = tmp2 = ''
            rst = False
            for prj in self.arr_prj:
                if ("IP" + pat.p_order) in prj.ip_list:
                    rst, tmp1, tmp2, tmp3 = self.get_prj_core(prj)
                    if rst:
                        ws.cell(row, column=9, value=tmp1)
                        ws.cell(row, column=10, value=tmp2)
                        break
            row = row + 1
        try:
            wb.save(self.file_pat_upload)
            self.textEdit.append(f'{self.file_pat_upload} 更新完成')
        except PermissionError:
            self.textEdit.append('写文件失败，关闭其他占用该文件的程序.' + self.file_pat_upload)
        finally:
            wb.close()

    def prepare_rdsummary_table(self):
        """
        企业研究开发活动汇总表
        """
        self.textEdit.setText('')
        self.update_data()
        filename = '2.企业研究开发活动汇总表（近三年执行的活动）.xlsx'
        file_rd_summary = self.workdir + '/' + filename
        if not os.path.exists(file_rd_summary):
            self.textEdit.append('文件不存在.' + filename)
            return

        wb = load_workbook(file_rd_summary)
        ws = wb.active
        ws.delete_rows(2, 200)
        row = 2
        for prj in self.arr_prj:
            ws.cell(row, column=1, value="RD" + prj.p_order)
            ws.cell(row, column=2, value=prj.p_name)
            # TODO   五、高技术服务/（四）高技术专业化服务
            # ws.cell(row, column=3, value=prj.p_class)
            # ws.cell(row, column=4, value=prj.p_got)
            # ws.cell(row, column=5, value=prj.p_no)
            ws.cell(row, column=6, value=prj.p_start)
            ws.cell(row, column=7, value=prj.p_end)
            ws.cell(row, column=8, value='企业自有技术')
            ws.cell(row, column=9, value=prj.ip_list)
            ws.cell(row, column=10, value=prj.p_money)
            ws.cell(row, column=15, value=prj.p_people)
            ws.cell(row, column=22, value=prj.p_product)
            rst, tmp1, tmp2, tmpcx  = self.get_prj_init(prj)
            if rst:
                ws.cell(row, column=16, value=tmp1)
                ws.cell(row, column=18, value=tmp2)
                ws.cell(row, column=21, value=tmpcx)

            rst, tmp1, tmp2, tmp = self.get_prj_core(prj)
            if rst:
                ws.cell(row, column=17, value=tmp)
                ws.cell(row, column=19, value=tmp1)
                ws.cell(row, column=20, value=tmp2)
            row = row + 1
        try:
            wb.save(file_rd_summary)
            self.textEdit.append(f'{filename} 更新完成')
        except PermissionError:
            self.textEdit.append('写文件失败，关闭其他占用该文件的程序.' + filename)
        finally:
            wb.close()


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())
